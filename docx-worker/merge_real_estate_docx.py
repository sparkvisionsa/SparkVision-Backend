#!/usr/bin/env python3
"""
دمج تقرير Word للتقارير العقارية (transactions) عبر المتغيرات المرئية
داخل قالب الشركة (report.docx) — نسخة مبسّطة مستوحاة من merge_docx.py
(الخاص بتقييم الآلات) لكن منفصلة تماماً عنه.

Usage: python merge_real_estate_docx.py < payload.json > output.docx
   or: python merge_real_estate_docx.py payload.json   (writes to payload.outputPath if given)

القالب حالياً يحتوي عدداً قليلاً من المتغيرات فقط (سيُوسَّع لاحقاً)؛
PLACEHOLDER_FIELDS هو المصدر الوحيد للربط بين اسم المتغير الظاهر في Word
ومفتاح القيمة القادم من الخادم — لا تُقرأ أي حقول أخرى لتحديد القيمة.
"""

from __future__ import annotations

import base64
import gc
import io
import json
import math
import os
import re
import sys
import traceback
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Callable, TypeVar, Union

from lxml import etree
from PIL import Image

ImageSource = Union[bytes, str]

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
XML_NS = "http://www.w3.org/XML/1998/namespace"

MERGE_PARTS_RE = re.compile(r"^word/(document|header\d+|footer\d+)\.xml$", re.I)
VISIBLE_VARIABLE_RE = re.compile(
    r"«(?P<guillemet>[^»\r\n]+)»|<<(?P<ascii>[^<>\r\n]+)>>"
)

# ─── قالب التقارير العقارية: أسماء المتغيرات المرئية → مفاتيح القيمة ───────────
# هذه هي القائمة الحالية فقط؛ القالب سيُوسَّع لاحقاً بمزيد من المتغيرات ويكفي
# إضافة سطر هنا عند ذلك — لا حاجة لتغيير منطق الاستبدال.
PLACEHOLDER_FIELDS: dict[str, str] = {
    "اسم العميل": "clientName",
    "النشاط العقاري": "clientActivity",
    "نوع الكيان القانوني": "legalEntityType",
    "نوع العقار": "propertyType",
    "المدينة": "city",
    "رقم التقرير": "reportNumber",
    "تاريخ التقييم": "valuationDate",
    "الحي": "neighborhood",
    "نوع الملكية": "ownershipType",
    "أساس التقييم": "valuationBasis",
}

# العنوان الذي تُدرج صور المعاملة بعده — قابل للتغيير عبر imageLayout.sectionHeading
DEFAULT_IMAGE_SECTION_HEADING = "الصور والملحقات"

DEFAULT_IMAGES_PER_ROW = 3
EMU_PER_INCH = 914400
DEFAULT_PAGE_WIDTH_EMU = int(8.27 * EMU_PER_INCH)
DEFAULT_PAGE_HEIGHT_EMU = int(11.69 * EMU_PER_INCH)
DEFAULT_PAGE_MARGIN_EMU = int(0.5 * EMU_PER_INCH)
IMAGE_CONTENT_WIDTH_RATIO = 0.95
IMAGE_GAP_PX = 3
PIXEL_EMU = int(EMU_PER_INCH / 96)
IMAGE_GAP_EMU = PIXEL_EMU * IMAGE_GAP_PX
IMAGE_GAP_DXA = 15 * IMAGE_GAP_PX
ASSET_IMAGE_MAX_SQUARE_PX = 1000
DOCUMENT_IMAGE_JPEG_QUALITY = 88
MAX_DRAWING_HEIGHT_EMU = int(20 * EMU_PER_INCH)


def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def log(msg: str) -> None:
    try:
        sys.stderr.buffer.write((msg + "\n").encode("utf-8", errors="replace"))
        sys.stderr.buffer.flush()
    except Exception:
        print(msg, file=sys.stderr)


def log_exception(prefix: str, exc: BaseException) -> None:
    log(f"{prefix}: {type(exc).__name__}: {exc!r}")


def sanitize_xml_text(value: str, strip: bool = True) -> str:
    if not value:
        return ""
    cleaned = re.sub(r"[\uD800-\uDFFF]", "", str(value))
    cleaned = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\uFFFE\uFFFF]", "", cleaned)
    return cleaned.strip() if strip else cleaned


def set_text_preserve_space(node: etree._Element, text: str) -> None:
    node.text = sanitize_xml_text(text, strip=False)
    if node.text and (node.text[:1].isspace() or node.text[-1:].isspace()):
        node.set(f"{{{XML_NS}}}space", "preserve")
    else:
        node.attrib.pop(f"{{{XML_NS}}}space", None)


def serialize_word_xml(root: etree._Element) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def find_ancestor(el: etree._Element | None, tag: str) -> etree._Element | None:
    while el is not None:
        if el.tag == tag:
            return el
        el = el.getparent()
    return None


def text_node_run(node: etree._Element) -> etree._Element | None:
    return find_ancestor(node, w("r"))


def paragraph_has_nested_story(para: etree._Element) -> bool:
    return any(candidate is not para for candidate in para.iter(w("p")))


def _paragraph_own_text_nodes_with_offsets(
    para: etree._Element,
) -> tuple[list[tuple[etree._Element, int, int]], str]:
    nodes: list[tuple[etree._Element, int, int]] = []
    parts: list[str] = []
    offset = 0
    for node in para.iter(w("t")):
        owning = node.getparent()
        while owning is not None and owning.tag != w("p"):
            owning = owning.getparent()
        if owning is not para:
            continue
        text = sanitize_xml_text(node.text or "", strip=False)
        node.text = text
        start = offset
        offset += len(text)
        nodes.append((node, start, offset))
        parts.append(text)
    return nodes, "".join(parts)


def normalize_placeholder_name(value: str) -> str:
    cleaned = re.sub(r"[\u200e\u200f\u202a-\u202e]", "", sanitize_xml_text(value or ""))
    cleaned = cleaned.strip()
    if cleaned.startswith("«") and cleaned.endswith("»"):
        cleaned = cleaned[1:-1]
    elif cleaned.startswith("<<") and cleaned.endswith(">>"):
        cleaned = cleaned[2:-2]
    cleaned = cleaned.replace("*", "_")
    return re.sub(r"\s+", "_", cleaned)


def placeholder_field_key(name: str) -> str | None:
    normalized = normalize_placeholder_name(name)
    for placeholder, field in PLACEHOLDER_FIELDS.items():
        if normalize_placeholder_name(placeholder) == normalized:
            return field
    return None


def placeholder_value(name: str, text_values: dict[str, str]) -> tuple[bool, str]:
    field = placeholder_field_key(name)
    if field is None:
        return False, ""
    raw = str(text_values.get(field, "") or "").strip()
    return True, sanitize_xml_text(raw, strip=False)


def visible_variable_name(match: re.Match[str]) -> str:
    return sanitize_xml_text(match.group("guillemet") or match.group("ascii") or "", strip=False)


def visible_variable_inner_span(match: re.Match[str]) -> tuple[int, int]:
    group_name = "guillemet" if match.group("guillemet") is not None else "ascii"
    return match.start(group_name), match.end(group_name)


def variable_style_score(text: str) -> int:
    return len(re.sub(r"[\W_]+", "", text, flags=re.UNICODE))


def select_visible_variable_text_node(
    nodes: list[tuple[etree._Element, int, int]],
    match: re.Match[str],
) -> etree._Element | None:
    inner_start, inner_end = visible_variable_inner_span(match)
    candidates: list[tuple[int, int, int, int, etree._Element]] = []
    for idx, (node, node_start, node_end) in enumerate(nodes):
        overlap_start = max(inner_start, node_start)
        overlap_end = min(inner_end, node_end)
        if overlap_end <= overlap_start:
            continue
        text = node.text or ""
        segment = text[overlap_start - node_start : overlap_end - node_start]
        run = text_node_run(node)
        has_rpr = int(run is not None and run.find(w("rPr")) is not None)
        candidates.append((variable_style_score(segment), has_rpr, len(segment), -idx, node))
    if candidates:
        return max(candidates, key=lambda item: item[:4])[4]
    for node, node_start, node_end in nodes:
        if node_end > match.start() and node_start < match.end():
            return node
    return None


def replace_text_range_in_selected_node(
    nodes: list[tuple[etree._Element, int, int]],
    start: int,
    end: int,
    replacement: str,
    target_node: etree._Element,
) -> bool:
    if start < 0 or end < start or not nodes:
        return False
    safe = sanitize_xml_text(replacement, strip=False)
    touched = False
    inserted = False
    for node, node_start, node_end in nodes:
        if node_end <= start or node_start >= end:
            continue
        text = node.text or ""
        prefix = text[: start - node_start] if node_start <= start < node_end else ""
        suffix = text[end - node_start :] if node_start < end <= node_end else ""
        value = safe if node is target_node else ""
        if node is target_node:
            inserted = True
        set_text_preserve_space(node, prefix + value + suffix)
        touched = True
    return touched and inserted


def replace_visible_variables(root: etree._Element, text_values: dict[str, str]) -> tuple[int, int]:
    found = 0
    filled = 0
    for para in root.iter(w("p")):
        nodes, full_text = _paragraph_own_text_nodes_with_offsets(para)
        if not nodes:
            continue
        matches = list(VISIBLE_VARIABLE_RE.finditer(full_text))
        found += len(matches)
        for match in reversed(matches):
            known, value = placeholder_value(visible_variable_name(match), text_values)
            if not known:
                continue
            target_node = select_visible_variable_text_node(nodes, match)
            if target_node is None:
                continue
            if replace_text_range_in_selected_node(nodes, match.start(), match.end(), value, target_node):
                if value.strip():
                    filled += 1
                nodes, full_text = _paragraph_own_text_nodes_with_offsets(para)
    return found, filled


def element_field_char_types(element: etree._Element) -> list[str]:
    return [str(node.get(w("fldCharType")) or "") for node in element.iter(w("fldChar"))]


def find_complex_field_end(children: list[etree._Element], start_idx: int) -> int | None:
    depth = 0
    saw_begin = False
    for idx in range(start_idx, len(children)):
        for field_type in element_field_char_types(children[idx]):
            if field_type == "begin":
                depth += 1
                saw_begin = True
            elif field_type == "end" and saw_begin:
                depth -= 1
                if depth == 0:
                    return idx
    return None


def field_result_elements(elements: list[etree._Element]) -> list[etree._Element]:
    separate_idx = None
    end_idx = None
    depth = 0
    for idx, element in enumerate(elements):
        for field_type in element_field_char_types(element):
            if field_type == "begin":
                depth += 1
            elif field_type == "separate" and depth == 1 and separate_idx is None:
                separate_idx = idx
            elif field_type == "end":
                depth = max(0, depth - 1)
                if depth == 0:
                    end_idx = idx
                    break
        if end_idx is not None:
            break
    if separate_idx is None or end_idx is None or end_idx <= separate_idx:
        return []
    return elements[separate_idx + 1 : end_idx]


def run_has_meaningful_content(run: etree._Element) -> bool:
    for child in run:
        if child.tag == w("rPr"):
            continue
        if child.tag == w("t") and (child.text or "").strip():
            return True
        if child.tag in (w("drawing"), w("pict"), w("tab"), w("br"), w("fldChar"), w("instrText")):
            return True
        if child.tag in (w("bookmarkStart"), w("bookmarkEnd")):
            continue
        return True
    return False


def cleanup_empty_runs(para: etree._Element) -> None:
    for run in list(para.findall(w("r"))):
        if not run_has_meaningful_content(run):
            has_bookmark = any(c.tag in (w("bookmarkStart"), w("bookmarkEnd")) for c in run)
            if not has_bookmark:
                parent = run.getparent()
                if parent is not None:
                    parent.remove(run)


def flatten_mail_merge_fields(root: etree._Element) -> int:
    flattened = 0
    for para in list(root.iter(w("p"))):
        if paragraph_has_nested_story(para):
            continue
        idx = 0
        while True:
            children = list(para)
            if idx >= len(children):
                break
            if "begin" not in element_field_char_types(children[idx]):
                idx += 1
                continue
            end_idx = find_complex_field_end(children, idx)
            if end_idx is None:
                idx += 1
                continue
            sequence = children[idx : end_idx + 1]
            instruction = "".join(
                node.text or "" for element in sequence for node in element.iter(w("instrText"))
            )
            if re.search(r"\bMERGEFIELD\b", instruction, re.IGNORECASE) is None:
                idx = end_idx + 1
                continue
            if re.search(r"\bPAGEREF\b|\bHYPERLINK\b|\bTOC\b", instruction, re.IGNORECASE):
                idx = end_idx + 1
                continue
            result_ids = {id(element) for element in field_result_elements(sequence)}
            kept = 0
            for element in sequence:
                if element.getparent() is not para:
                    continue
                if id(element) in result_ids or element.tag in (w("bookmarkStart"), w("bookmarkEnd")):
                    kept += 1
                    continue
                para.remove(element)
            flattened += 1
            idx += kept
        cleanup_empty_runs(para)
    return flattened


def apply_visible_variable_values(
    xml_bytes: bytes,
    text_values: dict[str, str] | None = None,
) -> tuple[bytes, int, int]:
    root = etree.fromstring(xml_bytes)
    clean_text_values = text_values or {}
    found, filled = replace_visible_variables(root, clean_text_values)
    flatten_mail_merge_fields(root)
    return serialize_word_xml(root), found, filled


def collect_template_placeholder_names(xml_bytes: bytes) -> list[str]:
    root = etree.fromstring(xml_bytes)
    names: list[str] = []
    for para in root.iter(w("p")):
        if paragraph_has_nested_story(para):
            continue
        visible = "".join(node.text or "" for node in para.iter(w("t")))
        for match in VISIBLE_VARIABLE_RE.finditer(visible):
            name = visible_variable_name(match)
            if name and name not in names:
                names.append(name)
    return names


# ─── تجهيز الصور ────────────────────────────────────────────────────────────

_IMAGE_PREP_T = TypeVar("_IMAGE_PREP_T")


def image_prep_worker_count(job_count: int = 1) -> int:
    if job_count <= 1:
        return 1
    cpus = os.cpu_count() or 2
    return max(1, min(8, cpus * 2, job_count))


def map_image_prep(items: list[Any], worker: Callable[[Any], _IMAGE_PREP_T]) -> list[_IMAGE_PREP_T | None]:
    if not items:
        return []
    workers = image_prep_worker_count(len(items))
    if workers == 1:
        out: list[_IMAGE_PREP_T | None] = []
        for item in items:
            try:
                out.append(worker(item))
            except Exception as exc:
                log_exception("image prep skipped", exc)
                out.append(None)
        return out
    results: list[_IMAGE_PREP_T | None] = [None] * len(items)
    with ThreadPoolExecutor(max_workers=workers) as pool:
        futures = {pool.submit(worker, item): idx for idx, item in enumerate(items)}
        for fut in as_completed(futures):
            idx = futures[fut]
            try:
                results[idx] = fut.result()
            except Exception as exc:
                log_exception("image prep skipped", exc)
                results[idx] = None
    return results


def resolve_image_bytes(source: ImageSource) -> bytes:
    if isinstance(source, (bytes, bytearray, memoryview)):
        return bytes(source)
    with open(str(source), "rb") as fh:
        return fh.read()


def collect_image_sources(payload: dict[str, Any], paths_key: str, b64_key: str) -> list[ImageSource]:
    paths = payload.get(paths_key) or []
    if isinstance(paths, list) and paths:
        out: list[ImageSource] = []
        for item in paths:
            if isinstance(item, str) and item.strip() and os.path.isfile(item):
                out.append(item)
        if out:
            return out
    encoded = payload.get(b64_key) or []
    if not isinstance(encoded, list):
        return []
    out_b64: list[ImageSource] = []
    for item in encoded:
        try:
            raw = base64.b64decode(item)
            if raw:
                out_b64.append(raw)
        except Exception:
            continue
    return out_b64


def jpeg_is_docx_safe_baseline(data: bytes) -> bool:
    if not data or len(data) < 4 or data[0:2] != b"\xff\xd8":
        return False
    i = 2
    found_sof0 = False
    found_identity_app = False
    length = len(data)
    while i < length - 1:
        if data[i] != 0xFF:
            i += 1
            continue
        while i < length and data[i] == 0xFF:
            i += 1
        if i >= length:
            break
        marker = data[i]
        i += 1
        if marker in (0xD8, 0xD9):
            continue
        if marker == 0x01 or 0xD0 <= marker <= 0xD7:
            continue
        if i + 1 >= length:
            break
        seg_len = (data[i] << 8) | data[i + 1]
        if seg_len < 2 or i + seg_len > length:
            break
        payload = data[i + 2 : i + seg_len]
        if marker == 0xC0:
            found_sof0 = True
        elif marker in (0xC1, 0xC2, 0xC3, 0xC5, 0xC6, 0xC7, 0xC9, 0xCA, 0xCB, 0xCD, 0xCE, 0xCF):
            return False
        elif marker == 0xE0 and payload.startswith(b"JFIF\x00"):
            found_identity_app = True
        elif marker == 0xE1 and payload.startswith(b"Exif\x00\x00"):
            found_identity_app = True
        i += seg_len
    return found_sof0 and found_identity_app


def _save_print_jpeg(img: Image.Image, quality: int = DOCUMENT_IMAGE_JPEG_QUALITY) -> bytes:
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    elif img.mode == "L":
        img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=max(60, min(100, int(quality))), optimize=False, progressive=False, subsampling=0)
    return out.getvalue()


def _canvas_pixel_size_for_cell(target_width_emu: int, target_height_emu: int, max_side_px: int) -> tuple[int, int]:
    tw = max(1, int(target_width_emu))
    th = max(1, int(target_height_emu))
    side = max(64, int(max_side_px))
    if tw >= th:
        return side, max(1, int(round(side * th / tw)))
    return max(1, int(round(side * tw / th))), side


def stretch_to_fill_canvas_jpeg_bytes(
    img_bytes: bytes,
    target_width_emu: int,
    target_height_emu: int,
    *,
    max_side_px: int = ASSET_IMAGE_MAX_SQUARE_PX,
    quality: int = DOCUMENT_IMAGE_JPEG_QUALITY,
) -> bytes:
    canvas_w, canvas_h = _canvas_pixel_size_for_cell(target_width_emu, target_height_emu, max_side_px)
    try:
        with Image.open(io.BytesIO(img_bytes)) as img:
            src_w, src_h = img.size
            if src_w <= 0 or src_h <= 0:
                return img_bytes
            if src_w == canvas_w and src_h == canvas_h and img.format == "JPEG" and jpeg_is_docx_safe_baseline(img_bytes):
                return img_bytes
            img.load()
            work = img.convert("RGB") if img.mode != "RGB" else img
            if work.size != (canvas_w, canvas_h):
                if work.width > canvas_w * 3 and work.height > canvas_h * 3:
                    work = work.resize((max(canvas_w, work.width // 2), max(canvas_h, work.height // 2)), Image.BILINEAR)
                work = work.resize((canvas_w, canvas_h), Image.LANCZOS)
            return _save_print_jpeg(work, quality)
    except Exception:
        return img_bytes


def prepare_image_png_fallback(img_bytes: bytes) -> bytes:
    img = Image.open(io.BytesIO(img_bytes))
    img.load()
    img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="PNG", optimize=False)
    return out.getvalue()


# ─── تخطيط صفحة python-docx ─────────────────────────────────────────────────

def docx_qn(tag: str):
    from docx.oxml.ns import qn
    return qn(tag)


def make_docx_element(tag: str):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)


def document_physical_page_box_emu(doc) -> tuple[int, int]:
    section = doc.sections[0]
    page_width = int(section.page_width or DEFAULT_PAGE_WIDTH_EMU)
    page_height = int(section.page_height or DEFAULT_PAGE_HEIGHT_EMU)
    return page_width or DEFAULT_PAGE_WIDTH_EMU, page_height or DEFAULT_PAGE_HEIGHT_EMU


def document_section_margins_emu(doc) -> tuple[int, int, int, int]:
    section = doc.sections[0]
    left = int(section.left_margin or DEFAULT_PAGE_MARGIN_EMU)
    right = int(section.right_margin or DEFAULT_PAGE_MARGIN_EMU)
    top = int(section.top_margin or DEFAULT_PAGE_MARGIN_EMU)
    bottom = int(section.bottom_margin or DEFAULT_PAGE_MARGIN_EMU)
    return left, right, top, bottom


def document_content_box_emu(doc) -> tuple[int, int]:
    page_width, page_height = document_physical_page_box_emu(doc)
    left, right, top, bottom = document_section_margins_emu(doc)
    content_width = max(1, page_width - left - right)
    content_height = max(1, page_height - top - bottom)
    return content_width, content_height


def set_docx_cell_margins(cell, margin_dxa: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(docx_qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = make_docx_element("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        node = make_docx_element(f"w:{side}")
        node.set(docx_qn("w:w"), str(margin_dxa))
        node.set(docx_qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_docx_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(docx_qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_borders = make_docx_element("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = make_docx_element(f"w:{border_name}")
        border.set(docx_qn("w:val"), "none")
        border.set(docx_qn("w:sz"), "0")
        border.set(docx_qn("w:space"), "0")
        border.set(docx_qn("w:color"), "auto")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def set_docx_table_cell_spacing(table, spacing_dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(docx_qn("w:tblCellSpacing"))
    if existing is not None:
        tbl_pr.remove(existing)
    spacing = make_docx_element("w:tblCellSpacing")
    spacing.set(docx_qn("w:w"), str(max(0, int(spacing_dxa))))
    spacing.set(docx_qn("w:type"), "dxa")
    tbl_pr.append(spacing)


def set_docx_table_width(table, width_emu: int) -> None:
    width_dxa = max(1, int(width_emu / EMU_PER_INCH * 1440))
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(docx_qn("w:tblW"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_w = make_docx_element("w:tblW")
    tbl_w.set(docx_qn("w:w"), str(width_dxa))
    tbl_w.set(docx_qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    layout = make_docx_element("w:tblLayout")
    layout.set(docx_qn("w:type"), "fixed")
    tbl_pr.append(layout)


def set_docx_table_indent(table, indent_emu: int) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(docx_qn("w:tblInd"))
    if existing is not None:
        tbl_pr.remove(existing)
    indent = make_docx_element("w:tblInd")
    indent.set(docx_qn("w:w"), str(max(0, int(indent_emu / EMU_PER_INCH * 1440))))
    indent.set(docx_qn("w:type"), "dxa")
    tbl_pr.append(indent)


def detach_docx_body_element(elem) -> Any:
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)
    return elem


def make_docx_page_break_element(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return detach_docx_body_element(p._element)


def block_text(block) -> str:
    return sanitize_xml_text("".join(node.text or "" for node in block.iter(docx_qn("w:t"))))


def normalize_heading_text(name: str) -> str:
    cleaned = re.sub(r"[\u200e\u200f\u202a-\u202e]", "", name or "")
    return re.sub(r"[\s_\-.،؛:\u060c\u061b\u0640]+", "", cleaned).strip().lower()


def find_body_heading_index(children: list[Any], heading: str) -> int | None:
    wanted = normalize_heading_text(heading)
    prefix_hit = None
    for idx in range(len(children) - 1, -1, -1):
        child = children[idx]
        if etree.QName(child).localname != "p":
            continue
        normalized = normalize_heading_text(block_text(child))
        if normalized == wanted:
            return idx
        if normalized.startswith(wanted) and prefix_hit is None:
            prefix_hit = idx
    return prefix_hit


def block_is_empty_insertion_spacer(block) -> bool:
    if etree.QName(block).localname != "p":
        return False
    if block_text(block):
        return False
    protected_tags = (docx_qn("w:br"), docx_qn("w:sectPr"), docx_qn("w:drawing"), docx_qn("w:pict"))
    return not any(any(True for _ in block.iter(tag)) for tag in protected_tags)


def remove_empty_spacers_after(body, heading_idx: int) -> int:
    insert_at = heading_idx + 1
    while True:
        children = list(body)
        if insert_at >= len(children) or not block_is_empty_insertion_spacer(children[insert_at]):
            return insert_at
        body.remove(children[insert_at])


def document_cell_dimensions_emu(
    doc,
    images_per_row: int,
    image_rows_per_page: int,
    *,
    content_width_ratio: float = IMAGE_CONTENT_WIDTH_RATIO,
    gap_emu: int = IMAGE_GAP_EMU,
) -> tuple[int, int]:
    _content_width, content_height = document_content_box_emu(doc)
    page_width, _page_height = document_physical_page_box_emu(doc)
    width_ratio = min(1.0, max(0.3, float(content_width_ratio)))
    gap = max(0, int(gap_emu))
    available_width = max(1, int(page_width * width_ratio))
    width_fit = max(1, (available_width - (images_per_row - 1) * gap) // images_per_row)
    height_fit = max(1, (content_height - (image_rows_per_page - 1) * gap) // image_rows_per_page)
    return width_fit, max(1, min(width_fit, height_fit))


def make_docx_image_table_element(
    doc,
    images: list[ImageSource],
    images_per_row: int,
    image_rows_per_page: int,
    *,
    content_width_ratio: float = IMAGE_CONTENT_WIDTH_RATIO,
    gap_dxa: int = IMAGE_GAP_DXA,
    gap_emu: int = IMAGE_GAP_EMU,
    asset_max_side_px: int = ASSET_IMAGE_MAX_SQUARE_PX,
) -> tuple[Any, int]:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu, Pt

    rows = max(1, math.ceil(len(images) / images_per_row))
    cell_width_emu, image_emu = document_cell_dimensions_emu(
        doc, images_per_row, image_rows_per_page,
        content_width_ratio=content_width_ratio, gap_emu=gap_emu,
    )
    cell_width_dxa = max(1, int(cell_width_emu / EMU_PER_INCH * 1440))
    width_ratio = min(1.0, max(0.3, float(content_width_ratio)))
    gap_dxa_val = max(0, int(gap_dxa))
    gap_emu_val = max(0, int(gap_emu))

    table = doc.add_table(rows=rows, cols=images_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_docx_table_borders_none(table)
    set_docx_table_cell_spacing(table, gap_dxa_val)
    set_docx_table_width(table, cell_width_emu * images_per_row + (images_per_row - 1) * gap_emu_val)

    page_width, _page_height = document_physical_page_box_emu(doc)
    left_margin, _right_margin, _top_margin, _bottom_margin = document_section_margins_emu(doc)
    physical_side_margin = max(0, int(page_width * (1 - width_ratio) / 2))
    set_docx_table_indent(table, physical_side_margin - left_margin)

    max_picture_height = min(image_emu, MAX_DRAWING_HEIGHT_EMU)

    def _prep(source: ImageSource) -> tuple[bytes, int, int]:
        raw = resolve_image_bytes(source)
        stretched = stretch_to_fill_canvas_jpeg_bytes(
            raw, cell_width_emu, image_emu, max_side_px=asset_max_side_px, quality=DOCUMENT_IMAGE_JPEG_QUALITY
        )
        return stretched, max(1, cell_width_emu), max(1, max_picture_height)

    prepared_images = map_image_prep(list(images), _prep)

    inserted = 0
    img_index = 0
    for row in table.rows:
        row.height = Emu(image_emu)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = Emu(cell_width_emu)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(docx_qn("w:tcW"))
            if tc_w is not None:
                tc_pr.remove(tc_w)
            tc_w = make_docx_element("w:tcW")
            tc_w.set(docx_qn("w:w"), str(cell_width_dxa))
            tc_w.set(docx_qn("w:type"), "dxa")
            tc_pr.append(tc_w)
            set_docx_cell_margins(cell, 0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            if img_index >= len(prepared_images):
                continue
            prepared = prepared_images[img_index]
            img_index += 1
            if prepared is None:
                continue
            img_bytes, pic_w, pic_h = prepared
            try:
                try:
                    para.add_run().add_picture(io.BytesIO(img_bytes), width=Emu(pic_w), height=Emu(pic_h))
                except Exception as jpeg_exc:
                    log_exception("image JPEG insert failed, trying PNG", jpeg_exc)
                    for run in list(para.runs):
                        el = run._element
                        parent = el.getparent()
                        if parent is not None:
                            parent.remove(el)
                    png_bytes = prepare_image_png_fallback(img_bytes)
                    para.add_run().add_picture(io.BytesIO(png_bytes), width=Emu(pic_w), height=Emu(pic_h))
                inserted += 1
            except Exception as exc:
                log_exception("image insert skipped", exc)
            finally:
                prepared_images[img_index - 1] = None

    return detach_docx_body_element(table._tbl), inserted


def insert_images_after_heading(
    doc,
    heading: str,
    images: list[ImageSource],
    images_per_row: int,
    images_per_page: int,
    asset_max_side_px: int,
) -> int:
    body = doc.element.body
    children = list(body)
    if not images:
        log("image section: no images to insert")
        return 0
    target_idx = find_body_heading_index(children, heading)
    if target_idx is None:
        log(f"image section heading {heading!r} was not found in document body")
        return 0
    log(f"image section heading found: {heading!r}")
    insert_at = remove_empty_spacers_after(body, target_idx)

    inserted = 0
    image_rows_per_page = max(1, math.ceil(images_per_page / images_per_row))
    page_number = 0
    for page_idx in range(0, len(images), images_per_page):
        if page_idx > 0:
            body.insert(insert_at, make_docx_page_break_element(doc))
            insert_at += 1
        table_elem, count = make_docx_image_table_element(
            doc,
            images[page_idx : page_idx + images_per_page],
            images_per_row,
            image_rows_per_page,
            asset_max_side_px=asset_max_side_px,
        )
        body.insert(insert_at, table_elem)
        insert_at += 1
        inserted += count
        page_number += 1
        if page_number % 10 == 0:
            gc.collect()
    return inserted


# ─── تطبيع معرّفات الرسم + التحقق من الحزمة ─────────────────────────────────

def _renumber_drawing_ids_in_xml(raw: bytes, next_id: int) -> tuple[bytes | None, int]:
    try:
        root = etree.fromstring(raw)
    except Exception:
        return None, next_id
    changed = False
    for el in root.iter():
        local = etree.QName(el).localname
        if local not in ("docPr", "cNvPr") or el.get("id") is None:
            continue
        el.set("id", str(next_id))
        next_id += 1
        changed = True
    if not changed:
        return None, next_id
    return serialize_word_xml(root), next_id


def normalize_docx_drawing_ids(docx_bytes: bytes) -> bytes:
    modified: dict[str, bytes] = {}
    next_id = 1
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        for name in zin.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            rewritten, next_id = _renumber_drawing_ids_in_xml(zin.read(name), next_id)
            if rewritten is not None:
                modified[name] = rewritten
        if not modified:
            return docx_bytes
        return write_docx_zip(zin, modified)


def write_docx_zip(zin: zipfile.ZipFile, modified: dict[str, bytes]) -> bytes:
    out_buf = io.BytesIO()
    names = zin.namelist()
    ordered: list[str] = []
    if "mimetype" in names:
        ordered.append("mimetype")
    for name in names:
        if name != "mimetype":
            ordered.append(name)
    for path in modified:
        if path not in ordered:
            ordered.append(path)
    with zipfile.ZipFile(out_buf, "w") as zout:
        for fname in ordered:
            if fname in modified:
                data = modified[fname]
            elif fname in names:
                data = zin.read(fname)
            else:
                continue
            compress = zipfile.ZIP_STORED if fname == "mimetype" or fname.startswith("word/media/") else zipfile.ZIP_DEFLATED
            zi = zipfile.ZipInfo(fname)
            zi.compress_type = compress
            zout.writestr(zi, data)
    return out_buf.getvalue()


def validate_docx_package(docx_bytes: bytes) -> None:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        bad = zf.testzip()
        if bad:
            raise ValueError(f"Corrupt zip member: {bad}")
        names = set(zf.namelist())
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise ValueError("Invalid docx package: required parts are missing")


def ensure_print_quality_settings(xml_bytes: bytes) -> bytes:
    root = etree.fromstring(xml_bytes)
    node = root.find(w("doNotAutoCompressPictures"))
    if node is None:
        node = etree.Element(w("doNotAutoCompressPictures"))
        root.append(node)
    node.set(w("val"), "true")
    return serialize_word_xml(root)


# ─── الدمج ─────────────────────────────────────────────────────────────────

def merge_package(payload: dict[str, Any]) -> bytes | None:
    global DOCUMENT_IMAGE_JPEG_QUALITY
    output_path = str(payload.get("outputPath") or "").strip() or None
    template_path = str(payload.get("templatePath") or "").strip()
    template_b64 = payload.get("templateBase64") or ""

    if template_path and os.path.isfile(template_path):
        with open(template_path, "rb") as fh:
            template_bytes = fh.read()
    elif template_b64:
        template_bytes = base64.b64decode(template_b64)
    else:
        raise ValueError("templatePath or templateBase64 missing")

    text_values = payload.get("textValues") or {}
    images = collect_image_sources(payload, "imagePaths", "imagesBase64")

    image_layout = payload.get("imageLayout") if isinstance(payload.get("imageLayout"), dict) else {}
    try:
        images_per_row = max(1, min(6, int(image_layout.get("imagesPerRow", DEFAULT_IMAGES_PER_ROW))))
    except (TypeError, ValueError):
        images_per_row = DEFAULT_IMAGES_PER_ROW
    default_images_per_page = images_per_row * 4
    try:
        images_per_page = max(images_per_row, int(image_layout.get("imagesPerPage", default_images_per_page)))
    except (TypeError, ValueError):
        images_per_page = default_images_per_page
    section_heading = str(image_layout.get("sectionHeading") or DEFAULT_IMAGE_SECTION_HEADING).strip()
    try:
        image_quality = max(60, min(100, int(image_layout.get("imageQuality", DOCUMENT_IMAGE_JPEG_QUALITY))))
    except (TypeError, ValueError):
        image_quality = DOCUMENT_IMAGE_JPEG_QUALITY

    DOCUMENT_IMAGE_JPEG_QUALITY = image_quality

    log(f"merge start: images={len(images)} perRow={images_per_row} perPage={images_per_page} output={'disk' if output_path else 'stdout'}")

    in_buf = io.BytesIO(template_bytes)
    del template_bytes
    modified: dict[str, bytes] = {}
    variable_names_found: list[str] = []
    variables_occurrences_found = 0
    variables_filled = 0

    with zipfile.ZipFile(in_buf, "r") as zin:
        names = zin.namelist()
        for fname in names:
            if MERGE_PARTS_RE.match(fname):
                raw = zin.read(fname)
                for variable_name in collect_template_placeholder_names(raw):
                    if variable_name not in variable_names_found:
                        variable_names_found.append(variable_name)
                updated, found, filled = apply_visible_variable_values(raw, text_values)
                modified[fname] = updated
                variables_occurrences_found += found
                variables_filled += filled

        if "word/settings.xml" in names:
            modified["word/settings.xml"] = ensure_print_quality_settings(
                modified.get("word/settings.xml", zin.read("word/settings.xml"))
            )

        result = write_docx_zip(zin, modified)

    del in_buf
    gc.collect()

    images_inserted = 0
    if images:
        from docx import Document

        doc = Document(io.BytesIO(result))
        images_inserted = insert_images_after_heading(
            doc, section_heading, images, images_per_row, images_per_page, ASSET_IMAGE_MAX_SQUARE_PX
        )
        images.clear()
        gc.collect()
        if output_path:
            doc.save(output_path)
            del doc
            gc.collect()
        else:
            out = io.BytesIO()
            doc.save(out)
            del doc
            result = out.getvalue()

    result = normalize_docx_drawing_ids(result) if not output_path or not images else result

    if output_path:
        if not images:
            with open(output_path, "wb") as fh:
                fh.write(result)
        else:
            # already written via doc.save(); still normalise drawing ids on disk
            with open(output_path, "rb") as fh:
                on_disk = fh.read()
            normalized = normalize_docx_drawing_ids(on_disk)
            with open(output_path, "wb") as fh:
                fh.write(normalized)
        with open(output_path, "rb") as fh:
            validate_docx_package(fh.read())
        result = None
    else:
        validate_docx_package(result)

    log(
        json.dumps(
            {
                "variablesFound": variable_names_found,
                "variablesOccurrencesFound": variables_occurrences_found,
                "variablesFilled": variables_filled,
                "imagesInserted": images_inserted,
            },
            ensure_ascii=False,
        )
    )
    return result


def main() -> None:
    try:
        if len(sys.argv) > 1:
            with open(sys.argv[1], encoding="utf-8") as fh:
                payload = json.load(fh)
        else:
            payload = json.load(sys.stdin)
        result = merge_package(payload)
        if result is not None:
            sys.stdout.buffer.write(result)
        else:
            sys.stdout.write("OK\n")
    except Exception:
        log(traceback.format_exc())
        sys.exit(1)


if __name__ == "__main__":
    main()
