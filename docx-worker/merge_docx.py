#!/usr/bin/env python3
"""
دمج تقرير Word عبر المتغيرات المرئية داخل قالب «تقرير تقييم.docx».
Usage: python merge_docx.py < payload.json > output.docx
"""

from __future__ import annotations

import base64
import gc
import io
import json
import math
import os
import posixpath
import re
import sys
import traceback
import zipfile
from copy import deepcopy
from typing import Any, Union

from lxml import etree
from PIL import Image

ImageSource = Union[bytes, str]

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
WP14_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
XML_NS = "http://www.w3.org/XML/1998/namespace"
IMAGE_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
)

# متغيرات قالب «تقرير تقييم.docx». المصدر الوحيد هو النص الظاهر بين « » أو << >>.
# القيم هي مفاتيح textValues القادمة من الخادم؛ لا تُقرأ أسماء حقول Word لتحديد القيمة.
PLACEHOLDER_FIELDS: dict[str, str] = {
    "عنوان_التقرير": "reportTitle",
    "العميل": "clientName",
    "تاريخ_إصدار_التقرير": "reportIssueDate",
    "الرقم_المرجعي": "reportReference",
    "اسلوب_التقييم": "valuationMethod",
    "أسلوب_التقييم": "valuationMethod",
    "الأسلوب_المستخدم": "valuationMethod",
    "الغرض_من_التقييم": "valuationPurpose",
    "اساس_القيمة": "valuationBasis",
    "تاريخ_التقييم": "valuationDate",
    "تاريخ_الاتفاقية": "agreementDate",
    "تاريخ_المعاينة": "inspectionDate",
    "أصلأصول": "assetSingularPlural",
    "نشاط_الشركة": "clientActivity",
    "ممثل_العميل": "clientRepresentativeName",
    "صفتة": "clientRepresentativeRole",
    "هوية_المستخدمين_الأخرين": "intendedUsers",
    "الأصل_المعنية_الأصل_محل_التقييم": "assetSubjectDescription",
    "أساس_القيمة_المستخدم": "valuationBasisDefinition",
    "فرضية_القيمة": "valuePremiseDefinition",
    "المدينة": "inspectionLocation",
    "رابط_قوقل_ماب": "inspectionMapUrl",
    "رأي_القيمة_رقما_وكتابتا": "finalValueOpinion",
}

IMAGE_SECTION_HEADINGS: dict[str, tuple[str, ...]] = {
    "valuation": (
        "مرفق 1: الوصف الجزئي وحسابات القيمة",
        "مرفق1: الوصف الجزئي وحسابات القيمة",
    ),
    "asset": (
        "مرفق 2: الصور الفوتوغرافية",
        "مرفق2: الصور الفوتوغرافية",
    ),
    "client": (
        "مرفق 3: المستندات المستلمة من العميل",
        "مرفق3: المستندات المستلمة من العميل",
    ),
}

REPORT_PREPARER_TABLE_HEADERS = (
    "بيانات المقيم",
    "دور المقيم",
    "التوقيع",
)
REPORT_PREPARER_DEFAULT_ROLES = (
    "الإدارة التنفيذية وتعميد ومراجعة المخرجات النهائية",
    "إعداد التقرير",
    "المعاينة",
)
REPORT_PREPARER_MAX_ROWS = 100
REPORT_SIGNATURE_MAX_SOURCE_BYTES = 20 * 1024 * 1024
# لوحة PNG بنسبة صورة القالب (~1.743) بدقة أعلى لطباعة أوضح بعد تكبير العرض.
REPORT_SIGNATURE_CANVAS_SIZE = (1400, 803)
# تكبير مساحة التوقيع الظاهرة في Word أمام كل مُعدّ (بالنسبة لأبعاد wp:extent في القالب).
REPORT_SIGNATURE_DISPLAY_SCALE = 1.85

# خط التقرير بعد الدمج: Tajawal لكل النص، مع الإبقاء على Cocon أسفل الغلاف فقط.
REPORT_BODY_FONT = "Tajawal"
REPORT_COVER_FOOTER_FONT = "CoconNextArabic-Light"
REPORT_COCON_FONT_NAMES = {
    "CoconNextArabic-Light",
    "Cocon® Next Arabic",
    "Cocon Next Arabic",
}
REPORT_COVER_FOOTER_MARKERS = (
    "الرقم المرجعي",
    "الرقم_المرجعي",
    "تاريخ التقرير",
    "تاريخ_إصدار_التقرير",
    "تاريخ إصدار التقرير",
)

VISIBLE_VARIABLE_RE = re.compile(
    r"«(?P<guillemet>[^»\r\n]+)»|<<(?P<ascii>[^<>\r\n]+)>>"
)
CLIENT_DOCS_IMAGES_PER_ROW = 2
CLIENT_DOCS_IMAGES_PER_PAGE = 4
# مرفق 3: ملء صفحة عمودية (طولية) مثل القالب الأساسي — ليس عرضية.
CLIENT_DOCS_CONTENT_WIDTH_RATIO = 0.95
CLIENT_DOCS_CONTENT_HEIGHT_RATIO = 0.92
CLIENT_DOCS_BOTTOM_MARGIN_PX = 0
CLIENT_DOCS_GAP_PX = 1
CLIENT_DOCS_TITLE_RESERVE_PX = 28
CLIENT_DOCS_SECTION_TITLE = "مرفق 3: المستندات المستلمة من العميل"

MERGE_PARTS_RE = re.compile(r"^word/(document|header\d+|footer\d+)\.xml$", re.I)
ASSET_IMAGE_PAGE_TITLE = "مرفق 2: الصور الفوتوغرافية"
IMAGE_PAGE_TITLE = ASSET_IMAGE_PAGE_TITLE
IMAGES_PER_ROW = 4
IMAGES_PER_PAGE = 20
IMAGE_ROWS_PER_PAGE = math.ceil(IMAGES_PER_PAGE / IMAGES_PER_ROW)
EMU_PER_INCH = 914400
PIXEL_DXA = 15
PIXEL_EMU = int(EMU_PER_INCH / 96)
CLIENT_DOCS_BOTTOM_MARGIN_EMU = PIXEL_EMU * CLIENT_DOCS_BOTTOM_MARGIN_PX
CLIENT_DOCS_GAP_DXA = PIXEL_DXA * CLIENT_DOCS_GAP_PX
CLIENT_DOCS_GAP_EMU = PIXEL_EMU * CLIENT_DOCS_GAP_PX
CLIENT_DOCS_TITLE_RESERVE_EMU = PIXEL_EMU * CLIENT_DOCS_TITLE_RESERVE_PX
IMAGE_HORIZONTAL_MARGIN_PX = 3
IMAGE_HORIZONTAL_MARGIN_DXA = PIXEL_DXA * IMAGE_HORIZONTAL_MARGIN_PX
IMAGE_HORIZONTAL_MARGIN_EMU = PIXEL_EMU * IMAGE_HORIZONTAL_MARGIN_PX
IMAGE_CONTENT_WIDTH_RATIO = 0.95
ASSET_IMAGE_GAP_DXA = IMAGE_HORIZONTAL_MARGIN_DXA
ASSET_IMAGE_GAP_EMU = IMAGE_HORIZONTAL_MARGIN_EMU
# صور الأصول تُخفَّض مع كثرة العدد. حسابات القيمة تبقى بدقة طباعة عالية للزوم الواضح.
ASSET_IMAGE_MAX_SQUARE_PX = 1100
VALUATION_IMAGE_MAX_DIMENSION_PX = 5200
VALUATION_IMAGE_JPEG_QUALITY = 96
DOCUMENT_IMAGE_JPEG_QUALITY = 92
ASSET_IMAGE_JPEG_QUALITY = 80


def resolve_image_bytes(source: ImageSource) -> bytes:
    """حمّل بايتات صورة من مسار ملف أو من bytes جاهزة — بدون الإبقاء على كل الصور في الذاكرة."""
    if isinstance(source, (bytes, bytearray, memoryview)):
        return bytes(source)
    path = str(source)
    with open(path, "rb") as fh:
        return fh.read()


def collect_image_sources(payload: dict[str, Any], paths_key: str, b64_key: str) -> list[ImageSource]:
    paths = payload.get(paths_key) or []
    if isinstance(paths, list) and paths:
        out: list[ImageSource] = []
        for item in paths:
            if isinstance(item, str) and item.strip() and os.path.isfile(item):
                out.append(item)
        if out:
            return out
    encoded = payload.get(b64_key) or []
    if not isinstance(encoded, list):
        return []
    out_b64: list[ImageSource] = []
    for item in encoded:
        try:
            raw = base64.b64decode(item)
            if raw:
                out_b64.append(raw)
        except Exception:
            continue
    return out_b64


def adaptive_asset_max_side(count: int) -> int:
    if count <= 80:
        return ASSET_IMAGE_MAX_SQUARE_PX
    if count <= 250:
        return 900
    if count <= 800:
        return 780
    if count <= 2000:
        return 680
    return 600


def adaptive_asset_quality(count: int, requested_quality: int) -> int:
    """اعتبر اختيار المستخدم سقفاً، وخفّضه فقط للتقارير ذات الصور الكثيرة جداً."""
    if count <= 80:
        ceiling = 95
    elif count <= 250:
        ceiling = 90
    elif count <= 800:
        ceiling = 84
    elif count <= 2000:
        ceiling = 78
    else:
        ceiling = 72
    return min(requested_quality, ceiling)


DEFAULT_PAGE_WIDTH_EMU = int(8.27 * EMU_PER_INCH)
DEFAULT_PAGE_HEIGHT_EMU = int(11.69 * EMU_PER_INCH)
DEFAULT_PAGE_MARGIN_EMU = int(0.5 * EMU_PER_INCH)
ARABIC_RE = re.compile(r"[\u0600-\u06ff]")
MOJIBAKE_RE = re.compile(r"[ØÙÃÂÐÑ]")


def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def log(msg: str) -> None:
    try:
        sys.stderr.buffer.write((msg + "\n").encode("utf-8", errors="replace"))
        sys.stderr.buffer.flush()
    except Exception:
        print(msg, file=sys.stderr)


def log_exception(prefix: str, exc: BaseException) -> None:
    log(f"{prefix}: {type(exc).__name__}: {exc!r}")


# حد رسم Word العملي (~22 بوصة). صور Excel الطويلة كانت تتجاوزه فتفشل add_picture بصمت.
MAX_DRAWING_HEIGHT_EMU = int(20 * EMU_PER_INCH)


def ensure_jpeg_bytes(img_bytes: bytes, quality: int = DOCUMENT_IMAGE_JPEG_QUALITY) -> bytes:
    """أعد ترميز أي صورة إلى JPEG/RGB نظيف قبل الإدراج في Word."""
    try:
        with Image.open(io.BytesIO(img_bytes)) as img:
            img.load()
            rgb = img.convert("RGB") if img.mode != "RGB" else img
            return _save_print_jpeg(rgb, quality)
    except Exception as exc:
        log_exception("ensure_jpeg_bytes failed", exc)
        return img_bytes


def normalize_heading_text(name: str) -> str:
    cleaned = re.sub(r"[\u200e\u200f\u202a-\u202e]", "", repair_mojibake_text(name or ""))
    return re.sub(r"[\s_\-.،؛:\u060c\u061b\u0640]+", "", cleaned).strip().lower()


def arabic_score(value: str) -> int:
    return len(ARABIC_RE.findall(value or ""))


def mojibake_score(value: str) -> int:
    return len(MOJIBAKE_RE.findall(value or ""))


def repair_mojibake_text(value: str) -> str:
    """Repairs Arabic UTF-8 text that was decoded as cp1252/latin-1 before merge."""
    original = str(value or "")
    if not original or mojibake_score(original) == 0:
        return original

    candidates: list[str] = []
    for encoding in ("cp1252", "latin1"):
        try:
            candidates.append(original.encode(encoding).decode("utf-8"))
        except Exception:
            continue

    best = original
    best_key = (arabic_score(original), -mojibake_score(original))
    for candidate in candidates:
        key = (arabic_score(candidate), -mojibake_score(candidate))
        if key > best_key:
            best = candidate
            best_key = key
    return best


def sanitize_xml_text(value: str, strip: bool = True) -> str:
    """يزيل محارف UTF-16 surrogate ومحارف التحكم — lxml يرفضها."""
    if not value:
        return ""
    cleaned = re.sub(r"[\uD800-\uDFFF]", "", repair_mojibake_text(str(value)))
    cleaned = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\uFFFE\uFFFF]", "", cleaned)
    return cleaned.strip() if strip else cleaned


def serialize_word_xml(root: etree._Element) -> bytes:
    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )


def normalize_header_floating_wraps(xml_bytes: bytes) -> tuple[bytes, int]:
    """
    اجعل رسومات الرأس خلف النص بلا التفاف.

    القالب يحتوي صورة رأس بطول الصفحة مع wrapTight؛ Word كان يلف جداول الصور
    حولها فيسمح بصف واحد فقط لكل صفحة ويضاعف فواصل الصفحات.
    """
    root = etree.fromstring(xml_bytes)
    changed = 0
    wrap_tags = {
        f"{{{WP_NS}}}wrapTight",
        f"{{{WP_NS}}}wrapSquare",
        f"{{{WP_NS}}}wrapThrough",
        f"{{{WP_NS}}}wrapTopAndBottom",
    }
    for anchor in root.iter(f"{{{WP_NS}}}anchor"):
        wrapping_children = [
            child for child in list(anchor)
            if child.tag in wrap_tags
        ]
        if not wrapping_children:
            continue
        for child in wrapping_children:
            index = anchor.index(child)
            anchor.remove(child)
            anchor.insert(index, etree.Element(f"{{{WP_NS}}}wrapNone"))
        if anchor.get("behindDoc") != "1":
            anchor.set("behindDoc", "1")
        if anchor.get("allowOverlap") != "1":
            anchor.set("allowOverlap", "1")
        changed += 1
    return serialize_word_xml(root), changed


def collect_report_preparers(payload: dict[str, Any]) -> tuple[bool, list[dict[str, Any]]]:
    """
    اقرأ صفوف معدّي التقرير من الحمولة.

    وجود المفتاح له دلالة مستقلة عن محتواه: عند وجوده يجب حذف صفوف القالب
    الثابتة حتى لو كانت القائمة فارغة. غيابه وحده يحافظ على القالب للتوافق
    مع المستدعين القدامى.
    """
    if "reportPreparers" not in payload:
        return False, []
    raw_rows = payload.get("reportPreparers")
    if not isinstance(raw_rows, list):
        return True, []

    rows: list[dict[str, Any]] = []
    for item in raw_rows[:REPORT_PREPARER_MAX_ROWS]:
        if not isinstance(item, dict):
            continue

        def text_field(*keys: str, limit: int) -> str:
            for key in keys:
                value = item.get(key)
                if value is None:
                    continue
                cleaned = sanitize_xml_text(str(value))
                if cleaned:
                    return cleaned[:limit]
            return ""

        index = len(rows)
        report_role = text_field("reportRole", "preparerRole", limit=2_000)
        if not report_role and index < len(REPORT_PREPARER_DEFAULT_ROLES):
            report_role = REPORT_PREPARER_DEFAULT_ROLES[index]

        rows.append(
            {
                "userId": text_field("userId", "id", limit=200),
                "name": text_field(
                    "reportDisplayName",
                    "displayName",
                    "name",
                    limit=500,
                ),
                "jobTitle": text_field("jobTitle", "roleLabel", limit=1_000),
                "membershipNo": text_field(
                    "membershipNo",
                    "membershipNumber",
                    limit=200,
                ),
                "reportRole": report_role,
                "signaturePath": str(item.get("signaturePath") or "").strip(),
                "signatureImageDataUrl": str(
                    item.get("signatureImageDataUrl") or ""
                ).strip(),
                "signatureBase64": str(
                    item.get("signatureBase64")
                    or item.get("signatureImageBase64")
                    or ""
                ).strip(),
            }
        )
    return True, rows


def report_preparer_signature_bytes(row: dict[str, Any]) -> bytes | None:
    path = str(row.get("signaturePath") or "").strip()
    if path and os.path.isfile(path):
        try:
            if os.path.getsize(path) > REPORT_SIGNATURE_MAX_SOURCE_BYTES:
                log(f"report signature skipped (too large): {path}")
                return None
            with open(path, "rb") as fh:
                data = fh.read(REPORT_SIGNATURE_MAX_SOURCE_BYTES + 1)
            return data if 0 < len(data) <= REPORT_SIGNATURE_MAX_SOURCE_BYTES else None
        except OSError as exc:
            log_exception("report signature path read failed", exc)

    encoded = str(row.get("signatureImageDataUrl") or "").strip()
    if encoded:
        match = re.match(
            r"^data:image/[a-z0-9.+-]+;base64,(?P<data>.+)$",
            encoded,
            flags=re.I | re.S,
        )
        encoded = match.group("data") if match else ""
    if not encoded:
        encoded = str(row.get("signatureBase64") or "").strip()
    if not encoded:
        return None
    try:
        compact = re.sub(r"\s+", "", encoded)
        raw = base64.b64decode(compact, validate=True)
        if 0 < len(raw) <= REPORT_SIGNATURE_MAX_SOURCE_BYTES:
            return raw
    except Exception as exc:
        log_exception("report signature base64 decode failed", exc)
    return None


def prepare_report_signature_png(source_bytes: bytes) -> bytes | None:
    """
    حوّل التوقيع إلى PNG شفاف بنسبة صورة القالب نفسها.

    تثبيت نسبة اللوحة يحافظ على ملاءمة الصورة داخل wp:extent بعد التكبير.
    """
    if not source_bytes:
        return None
    try:
        with Image.open(io.BytesIO(source_bytes)) as source:
            width, height = source.size
            if width <= 0 or height <= 0 or width * height > 50_000_000:
                return None
            source.seek(0)
            source.load()
            rgba = source.convert("RGBA")

        alpha_bounds = rgba.getchannel("A").getbbox()
        if alpha_bounds is None:
            return None
        rgba = rgba.crop(alpha_bounds)

        canvas_width, canvas_height = REPORT_SIGNATURE_CANVAS_SIZE
        padding = max(8, int(min(canvas_width, canvas_height) * 0.035))
        max_width = max(1, canvas_width - padding * 2)
        max_height = max(1, canvas_height - padding * 2)
        resampling = getattr(Image, "Resampling", Image).LANCZOS
        rgba.thumbnail((max_width, max_height), resampling)

        canvas = Image.new("RGBA", (canvas_width, canvas_height), (255, 255, 255, 0))
        offset = (
            (canvas_width - rgba.width) // 2,
            (canvas_height - rgba.height) // 2,
        )
        canvas.alpha_composite(rgba, dest=offset)
        output = io.BytesIO()
        canvas.save(output, format="PNG", optimize=True)
        return output.getvalue()
    except Exception as exc:
        log_exception("report signature PNG preparation failed", exc)
        return None


def _scale_numeric_attr(element: etree._Element, attr_name: str, scale: float) -> None:
    raw = element.get(attr_name)
    if raw is None:
        # بعض القوالب تستخدم السمة بأسماء نطاق Word.
        namespaced = f"{{{W_NS}}}{attr_name}"
        raw = element.get(namespaced)
        attr_name = namespaced if raw is not None else attr_name
    if raw is None:
        return
    try:
        element.set(attr_name, str(max(1, int(round(int(raw) * scale)))))
    except ValueError:
        return


def scale_signature_drawing(drawing: etree._Element, scale: float) -> None:
    """كبّر إطار رسم التوقيع (wp:extent + a:ext) داخل خلية المقيم."""
    if abs(scale - 1.0) < 1e-6:
        return
    for extent in drawing.iter(f"{{{WP_NS}}}extent"):
        _scale_numeric_attr(extent, "cx", scale)
        _scale_numeric_attr(extent, "cy", scale)
    for ext in drawing.iter(f"{{{A_NS}}}ext"):
        _scale_numeric_attr(ext, "cx", scale)
        _scale_numeric_attr(ext, "cy", scale)


def scale_signature_cell_width(cell: etree._Element, scale: float) -> None:
    """وسّع عمود التوقيع لاستيعاب الصورة الأكبر."""
    if abs(scale - 1.0) < 1e-6:
        return
    for tc_pr in cell.findall(w("tcPr")):
        for tc_w in tc_pr.findall(w("tcW")):
            _scale_numeric_attr(tc_w, "w", scale)


def scale_preparer_signature_grid_column(table: etree._Element, scale: float) -> None:
    """وسّع عمود التوقيع في tblGrid (العمود الثالث)."""
    if abs(scale - 1.0) < 1e-6:
        return
    grid = table.find(w("tblGrid"))
    if grid is None:
        return
    cols = grid.findall(w("gridCol"))
    if len(cols) < 3:
        return
    _scale_numeric_attr(cols[2], "w", scale)


def _element_visible_text(element: etree._Element) -> str:
    return "".join(node.text or "" for node in element.iter(w("t")))


def find_report_preparer_table(root: etree._Element) -> etree._Element | None:
    expected = tuple(normalize_heading_text(value) for value in REPORT_PREPARER_TABLE_HEADERS)
    matches: list[etree._Element] = []
    for table in root.iter(w("tbl")):
        rows = table.findall(w("tr"))
        if not rows:
            continue
        cells = rows[0].findall(w("tc"))
        if len(cells) != len(expected):
            continue
        actual = tuple(
            normalize_heading_text(_element_visible_text(cell))
            for cell in cells
        )
        if actual == expected:
            matches.append(table)
    if len(matches) > 1:
        raise ValueError("Ambiguous report preparer table: multiple matching tables")
    return matches[0] if matches else None


def set_paragraph_text_preserving_format(
    paragraph: etree._Element,
    value: str,
) -> None:
    """غيّر w:t فقط واترك pPr/rPr وبقية بنية الفقرة كما هي."""
    clean_value = sanitize_xml_text(value, strip=False)
    text_nodes = list(paragraph.iter(w("t")))
    if not text_nodes:
        if not clean_value:
            return
        runs = paragraph.findall(w("r"))
        if runs:
            run = runs[0]
        else:
            run = etree.Element(w("r"))
            paragraph.append(run)
        text_node = etree.Element(w("t"))
        run_properties = run.find(w("rPr"))
        run.insert(1 if run_properties is not None else 0, text_node)
        text_nodes = [text_node]

    for index, node in enumerate(text_nodes):
        node.attrib.pop(f"{{{XML_NS}}}space", None)
        set_text_preserve_space(node, clean_value if index == 0 else "")


def _remove_drawings(container: etree._Element) -> None:
    for drawing in list(container.iter(w("drawing"))):
        parent = drawing.getparent()
        if parent is not None:
            parent.remove(drawing)


def _relationship_xml(root: etree._Element) -> bytes:
    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )


def _ensure_png_content_type(xml_bytes: bytes) -> bytes:
    root = etree.fromstring(xml_bytes)
    for child in root:
        if (
            etree.QName(child).localname == "Default"
            and (child.get("Extension") or "").lower() == "png"
        ):
            return xml_bytes
    node = etree.Element(f"{{{CONTENT_TYPES_NS}}}Default")
    node.set("Extension", "png")
    node.set("ContentType", "image/png")
    root.append(node)
    return _relationship_xml(root)


def _next_relationship_id(used_ids: set[str], index: int) -> str:
    suffix = max(1, index)
    while True:
        candidate = f"rIdReportPreparer{suffix}"
        if candidate not in used_ids:
            used_ids.add(candidate)
            return candidate
        suffix += 1


def _next_media_part_name(used_names: set[str], index: int) -> tuple[str, str]:
    suffix = max(1, index)
    while True:
        target = f"media/report-preparer-signature-{suffix}.png"
        package_name = f"word/{target}"
        if package_name not in used_names:
            used_names.add(package_name)
            return package_name, target
        suffix += 1


def _relationship_ids_used_in_document(root: etree._Element) -> set[str]:
    used: set[str] = set()
    for element in root.iter():
        for attribute_name, value in element.attrib.items():
            try:
                if etree.QName(attribute_name).namespace == R_NS and value:
                    used.add(value)
            except ValueError:
                continue
    return used


def inject_report_preparers(
    document_xml: bytes,
    document_rels_xml: bytes,
    content_types_xml: bytes,
    preparers: list[dict[str, Any]],
    package_names: set[str],
) -> tuple[bytes, bytes, bytes, dict[str, bytes], dict[str, int]]:
    """
    استبدل صفوف المقيمين الثابتة بصفوف ديناميكية مستنسخة من القالب.

    لا تُعاد كتابة خصائص الجدول/الخلايا/الفقرات/الخطوط. التغييرات الوحيدة
    داخل الصف المستنسخ هي النص الظاهر وعلاقة صورة التوقيع ومعرّفات الرسم.
    """
    root = etree.fromstring(document_xml)
    table = find_report_preparer_table(root)
    if table is None:
        raise ValueError(
            "Report preparer table not found (expected headers: "
            + " / ".join(REPORT_PREPARER_TABLE_HEADERS)
            + ")"
        )

    rows = table.findall(w("tr"))
    if not rows:
        raise ValueError("Report preparer table has no header row")
    header_row = rows[0]
    prototype = deepcopy(rows[1]) if len(rows) > 1 else None
    if preparers and prototype is None:
        raise ValueError("Report preparer table has no prototype data row")

    old_relationship_ids: set[str] = set()
    for old_row in rows[1:]:
        for blip in old_row.iter(f"{{{A_NS}}}blip"):
            relationship_id = blip.get(f"{{{R_NS}}}embed")
            if relationship_id:
                old_relationship_ids.add(relationship_id)
        table.remove(old_row)

    scale_preparer_signature_grid_column(table, REPORT_SIGNATURE_DISPLAY_SCALE)
    # وسّع خلية التوقيع في صف الرأس أيضاً لتبقى الأعمدة متناسقة.
    header_cells = header_row.findall(w("tc"))
    if len(header_cells) >= 3:
        scale_signature_cell_width(header_cells[2], REPORT_SIGNATURE_DISPLAY_SCALE)

    try:
        relationships_root = etree.fromstring(document_rels_xml)
    except Exception:
        relationships_root = etree.Element(f"{{{PACKAGE_REL_NS}}}Relationships")
    relationship_nodes = [
        node
        for node in relationships_root
        if etree.QName(node).localname == "Relationship"
    ]
    used_relationship_ids = {
        node.get("Id") or ""
        for node in relationship_nodes
        if node.get("Id")
    }

    max_drawing_id = 0
    for element in root.iter():
        if etree.QName(element).localname not in ("docPr", "cNvPr"):
            continue
        try:
            max_drawing_id = max(max_drawing_id, int(element.get("id") or "0"))
        except ValueError:
            continue
    used_anchor_ids = {
        str(anchor.get(f"{{{WP14_NS}}}anchorId") or "").upper()
        for anchor in root.iter(f"{{{WP_NS}}}anchor")
        if anchor.get(f"{{{WP14_NS}}}anchorId")
    }
    used_edit_ids = {
        str(anchor.get(f"{{{WP14_NS}}}editId") or "").upper()
        for anchor in root.iter(f"{{{WP_NS}}}anchor")
        if anchor.get(f"{{{WP14_NS}}}editId")
    }
    next_anchor_number = 0x71000000
    next_edit_number = 0x72000000

    def unique_hex_id(used: set[str], start: int) -> tuple[str, int]:
        number = start
        while True:
            candidate = f"{number & 0xFFFFFFFF:08X}"
            number += 1
            if candidate not in used:
                used.add(candidate)
                return candidate, number

    added_parts: dict[str, bytes] = {}
    signatures_inserted = 0
    for index, preparer in enumerate(preparers, start=1):
        assert prototype is not None
        row = deepcopy(prototype)
        cells = row.findall(w("tc"))
        if len(cells) != 3:
            raise ValueError("Report preparer prototype row must contain three cells")

        identity_paragraphs = cells[0].findall(w("p"))
        role_paragraphs = cells[1].findall(w("p"))
        if len(identity_paragraphs) < 3 or not role_paragraphs:
            raise ValueError("Report preparer prototype row has an unexpected structure")

        set_paragraph_text_preserving_format(
            identity_paragraphs[0],
            str(preparer.get("name") or ""),
        )
        set_paragraph_text_preserving_format(
            identity_paragraphs[1],
            str(preparer.get("jobTitle") or ""),
        )
        membership_no = str(preparer.get("membershipNo") or "").strip()
        set_paragraph_text_preserving_format(
            identity_paragraphs[2],
            f"عضوية رقم: {membership_no}" if membership_no else "",
        )
        for extra_paragraph in identity_paragraphs[3:]:
            set_paragraph_text_preserving_format(extra_paragraph, "")
        set_paragraph_text_preserving_format(
            role_paragraphs[0],
            str(preparer.get("reportRole") or ""),
        )
        for extra_paragraph in role_paragraphs[1:]:
            set_paragraph_text_preserving_format(extra_paragraph, "")

        signature_png = None
        signature_source = report_preparer_signature_bytes(preparer)
        if signature_source:
            signature_png = prepare_report_signature_png(signature_source)

        signature_cell = cells[2]
        drawings = list(signature_cell.iter(w("drawing")))
        if not signature_png or not drawings:
            _remove_drawings(signature_cell)
        else:
            drawing = drawings[0]
            for extra_drawing in drawings[1:]:
                parent = extra_drawing.getparent()
                if parent is not None:
                    parent.remove(extra_drawing)
            blips = list(drawing.iter(f"{{{A_NS}}}blip"))
            if not blips:
                _remove_drawings(signature_cell)
            else:
                package_name, relationship_target = _next_media_part_name(
                    package_names,
                    index,
                )
                relationship_id = _next_relationship_id(
                    used_relationship_ids,
                    index,
                )
                relationship = etree.Element(
                    f"{{{PACKAGE_REL_NS}}}Relationship"
                )
                relationship.set("Id", relationship_id)
                relationship.set("Type", IMAGE_REL_TYPE)
                relationship.set("Target", relationship_target)
                relationships_root.append(relationship)
                blips[0].set(f"{{{R_NS}}}embed", relationship_id)
                added_parts[package_name] = signature_png
                scale_signature_drawing(drawing, REPORT_SIGNATURE_DISPLAY_SCALE)

                max_drawing_id += 1
                for element in drawing.iter():
                    local_name = etree.QName(element).localname
                    if local_name in ("docPr", "cNvPr") and element.get("id") is not None:
                        element.set("id", str(max_drawing_id))
                        element.set("name", f"Report preparer signature {index}")
                for anchor in drawing.iter(f"{{{WP_NS}}}anchor"):
                    anchor_id, next_anchor_number = unique_hex_id(
                        used_anchor_ids,
                        next_anchor_number,
                    )
                    edit_id, next_edit_number = unique_hex_id(
                        used_edit_ids,
                        next_edit_number,
                    )
                    anchor.set(f"{{{WP14_NS}}}anchorId", anchor_id)
                    anchor.set(f"{{{WP14_NS}}}editId", edit_id)
                signatures_inserted += 1

        scale_signature_cell_width(signature_cell, REPORT_SIGNATURE_DISPLAY_SCALE)
        table.append(row)

    # لا تترك علاقات التوقيعات الثابتة فعّالة بعد حذف صفوفها. تبقى وسائطها
    # اليتيمة داخل الحزمة فقط إذا كانت جزءاً من القالب؛ وهي غير قابلة للعرض.
    used_after = _relationship_ids_used_in_document(root)
    for relationship in list(relationships_root):
        relationship_id = relationship.get("Id") or ""
        if relationship_id in old_relationship_ids and relationship_id not in used_after:
            relationships_root.remove(relationship)

    # تأكيد أن صف الرأس بقي أول عنصر صف داخل الجدول.
    if table.findall(w("tr"))[0] is not header_row:
        raise ValueError("Report preparer header row order changed unexpectedly")

    updated_document = serialize_word_xml(root)
    updated_relationships = _relationship_xml(relationships_root)
    updated_content_types = (
        _ensure_png_content_type(content_types_xml)
        if signatures_inserted
        else content_types_xml
    )
    return (
        updated_document,
        updated_relationships,
        updated_content_types,
        added_parts,
        {
            "tableFound": 1,
            "rowsRemoved": max(0, len(rows) - 1),
            "preparersInserted": len(preparers),
            "signaturesInserted": signatures_inserted,
        },
    )


def text_nodes_with_offsets(container: etree._Element) -> tuple[list[tuple[etree._Element, int, int]], str]:
    nodes: list[tuple[etree._Element, int, int]] = []
    parts: list[str] = []
    offset = 0
    for node in container.iter(w("t")):
        text = sanitize_xml_text(node.text or "", strip=False)
        node.text = text
        parts.append(text)
        start = offset
        offset += len(text)
        nodes.append((node, start, offset))
    return nodes, "".join(parts)


def set_text_preserve_space(node: etree._Element, text: str) -> None:
    node.text = sanitize_xml_text(text, strip=False)
    if node.text and (node.text[:1].isspace() or node.text[-1:].isspace()):
        node.set(f"{{{XML_NS}}}space", "preserve")


def normalize_placeholder_name(value: str) -> str:
    cleaned = re.sub(
        r"[\u200e\u200f\u202a-\u202e]",
        "",
        sanitize_xml_text(value or ""),
    )
    cleaned = cleaned.strip()
    if cleaned.startswith("«") and cleaned.endswith("»"):
        cleaned = cleaned[1:-1]
    elif cleaned.startswith("<<") and cleaned.endswith(">>"):
        cleaned = cleaned[2:-2]
    cleaned = cleaned.replace("*", "_")
    return re.sub(r"\s+", "_", cleaned)


def placeholder_field_key(name: str) -> str | None:
    normalized = normalize_placeholder_name(name)
    for placeholder, field in PLACEHOLDER_FIELDS.items():
        if normalize_placeholder_name(placeholder) == normalized:
            return field
    return None


def placeholder_value(name: str, text_values: dict[str, str]) -> tuple[bool, str]:
    field = placeholder_field_key(name)
    if field is None:
        return False, ""
    return True, sanitize_xml_text(str(text_values.get(field, "")), strip=False)


def paragraph_has_nested_story(para: etree._Element) -> bool:
    return any(candidate is not para for candidate in para.iter(w("p")))


def element_field_char_types(element: etree._Element) -> list[str]:
    return [
        str(node.get(w("fldCharType")) or "")
        for node in element.iter(w("fldChar"))
    ]


def field_result_elements(
    elements: list[etree._Element],
) -> list[etree._Element]:
    separate_idx: int | None = None
    end_idx: int | None = None
    depth = 0
    for idx, element in enumerate(elements):
        for field_type in element_field_char_types(element):
            if field_type == "begin":
                depth += 1
            elif field_type == "separate" and depth == 1 and separate_idx is None:
                separate_idx = idx
            elif field_type == "end":
                depth = max(0, depth - 1)
                if depth == 0:
                    end_idx = idx
                    break
        if end_idx is not None:
            break
    if separate_idx is None or end_idx is None or end_idx <= separate_idx:
        return []
    return elements[separate_idx + 1 : end_idx]


def find_complex_field_end(
    children: list[etree._Element],
    start_idx: int,
) -> int | None:
    depth = 0
    saw_begin = False
    for idx in range(start_idx, len(children)):
        for field_type in element_field_char_types(children[idx]):
            if field_type == "begin":
                depth += 1
                saw_begin = True
            elif field_type == "end" and saw_begin:
                depth -= 1
                if depth == 0:
                    return idx
    return None


def visible_variable_name(match: re.Match[str]) -> str:
    return sanitize_xml_text(
        match.group("guillemet") or match.group("ascii") or "",
        strip=False,
    )


def visible_variable_inner_span(match: re.Match[str]) -> tuple[int, int]:
    group_name = "guillemet" if match.group("guillemet") is not None else "ascii"
    return match.start(group_name), match.end(group_name)


def text_node_run(node: etree._Element) -> etree._Element | None:
    return find_ancestor(node, w("r"))


def variable_style_score(text: str) -> int:
    """قيّم محتوى اسم المتغير نفسه، مع تجاهل الأقواس والفراغات والشرطات السفلية."""
    return len(re.sub(r"[\W_]+", "", text, flags=re.UNICODE))


def select_visible_variable_text_node(
    nodes: list[tuple[etree._Element, int, int]],
    match: re.Match[str],
) -> etree._Element | None:
    """
    اختر run اسم المتغير لا run علامة الفتح.

    في القوالب التي توزّع «، الاسم، » على runs مختلفة يكون run الاسم هو صاحب
    الخط والحجم المقصودين. نختار أعلى عدد من حروف الاسم، ثم نفضّل run ذا rPr.
    """
    inner_start, inner_end = visible_variable_inner_span(match)
    candidates: list[tuple[int, int, int, int, etree._Element]] = []
    for idx, (node, node_start, node_end) in enumerate(nodes):
        overlap_start = max(inner_start, node_start)
        overlap_end = min(inner_end, node_end)
        if overlap_end <= overlap_start:
            continue
        text = node.text or ""
        segment = text[
            overlap_start - node_start : overlap_end - node_start
        ]
        run = text_node_run(node)
        has_rpr = int(run is not None and run.find(w("rPr")) is not None)
        candidates.append(
            (
                variable_style_score(segment),
                has_rpr,
                len(segment),
                -idx,
                node,
            )
        )
    if candidates:
        return max(candidates, key=lambda item: item[:4])[4]

    for node, node_start, node_end in nodes:
        if node_end > match.start() and node_start < match.end():
            return node
    return None


def replace_text_range_in_selected_node(
    nodes: list[tuple[etree._Element, int, int]],
    start: int,
    end: int,
    replacement: str,
    target_node: etree._Element,
) -> bool:
    """احذف كامل العلامة من كل runs وضع القيمة داخل run الاسم المختار فقط."""
    if start < 0 or end < start or not nodes:
        return False
    safe = sanitize_xml_text(replacement, strip=False)
    touched = False
    inserted = False

    for node, node_start, node_end in nodes:
        if node_end <= start or node_start >= end:
            continue
        text = node.text or ""
        prefix = text[: start - node_start] if node_start <= start < node_end else ""
        suffix = text[end - node_start :] if node_start < end <= node_end else ""
        value = safe if node is target_node else ""
        if node is target_node:
            inserted = True
        set_text_preserve_space(node, prefix + value + suffix)
        touched = True

    return touched and inserted


def replace_visible_variables(
    root: etree._Element,
    text_values: dict[str, str],
) -> tuple[int, int]:
    """
    استبدل حصرياً المتغيرات المرئية «name» أو <<name>>، حتى عند انقسامها عبر runs.

    لا تُقرأ تعليمات حقول Word لتحديد القيمة.
    """
    found = 0
    filled = 0
    for para in root.iter(w("p")):
        if paragraph_has_nested_story(para):
            continue
        nodes, full_text = text_nodes_with_offsets(para)
        matches = list(VISIBLE_VARIABLE_RE.finditer(full_text))
        found += len(matches)
        for match in reversed(matches):
            known, value = placeholder_value(
                visible_variable_name(match),
                text_values,
            )
            if not known:
                continue
            target_node = select_visible_variable_text_node(nodes, match)
            if target_node is None:
                continue
            if replace_text_range_in_selected_node(
                nodes,
                match.start(),
                match.end(),
                value,
                target_node,
            ):
                if value.strip():
                    filled += 1
                nodes, full_text = text_nodes_with_offsets(para)
    return found, filled


def flatten_mail_merge_fields(root: etree._Element) -> int:
    """
    سطّح كل MERGEFIELD بعد استبدال النص المرئي.

    نحتفظ بعناصر النتيجة الظاهرة كما هي ونحذف begin/instrText/separate/end؛
    اسم MERGEFIELD لا يشارك مطلقاً في اختيار قيمة المتغير.
    """
    flattened = 0
    for para in list(root.iter(w("p"))):
        if paragraph_has_nested_story(para):
            continue
        idx = 0
        while True:
            children = list(para)
            if idx >= len(children):
                break
            if "begin" not in element_field_char_types(children[idx]):
                idx += 1
                continue
            end_idx = find_complex_field_end(children, idx)
            if end_idx is None:
                idx += 1
                continue
            sequence = children[idx : end_idx + 1]
            instruction = "".join(
                node.text or ""
                for element in sequence
                for node in element.iter(w("instrText"))
            )
            if re.search(r"\bMERGEFIELD\b", instruction, re.IGNORECASE) is None:
                idx = end_idx + 1
                continue

            result_ids = {id(element) for element in field_result_elements(sequence)}
            kept = 0
            for element in sequence:
                if element.getparent() is not para:
                    continue
                if id(element) in result_ids:
                    kept += 1
                    continue
                para.remove(element)
            flattened += 1
            idx += kept
        cleanup_empty_runs(para)
    return flattened


def find_ancestor(el: etree._Element | None, tag: str) -> etree._Element | None:
    while el is not None:
        if el.tag == tag:
            return el
        el = el.getparent()
    return None
























































def run_has_meaningful_content(run: etree._Element) -> bool:
    for child in run:
        if child.tag == w("rPr"):
            continue
        if child.tag == w("t") and (child.text or "").strip():
            return True
        if child.tag in (w("drawing"), w("pict"), w("tab"), w("br"), w("fldChar"), w("instrText")):
            return True
        if child.tag in (w("bookmarkStart"), w("bookmarkEnd")):
            continue
        return True
    return False


def cleanup_empty_runs(para: etree._Element) -> None:
    for run in list(para.findall(w("r"))):
        if not run_has_meaningful_content(run):
            has_bookmark = any(
                c.tag in (w("bookmarkStart"), w("bookmarkEnd")) for c in run
            )
            if not has_bookmark:
                parent = run.getparent()
                if parent is not None:
                    parent.remove(run)






def validate_part_xml(xml_bytes: bytes, part_name: str) -> None:
    root = etree.fromstring(xml_bytes)

    def nested_run_is_embedded_story(nested: etree._Element, outer: etree._Element) -> bool:
        """Text boxes inside drawings store their own w:p/w:r tree inside an outer run."""
        saw_paragraph = False
        node = nested.getparent()
        while node is not None and node is not outer:
            if node.tag == w("txbxContent"):
                return True
            if node.tag == w("p"):
                saw_paragraph = True
            if saw_paragraph and node.tag in (w("drawing"), w("pict")):
                return True
            node = node.getparent()
        return False

    for run in root.iter(w("r")):
        for nested in run.iter(w("r")):
            if nested is not run:
                if nested_run_is_embedded_story(nested, run):
                    continue
                raise ValueError(f"Nested runs detected in {part_name}")
    for text in root.iter(w("t")):
        if text.getparent() is None or text.getparent().tag != w("r"):
            raise ValueError(f"Orphan w:t detected in {part_name}")
    for ppr in root.iter(w("pPr")):
        children = list(ppr)
        tags = [child.tag for child in children]
        if w("jc") in tags and w("bidi") in tags and tags.index(w("jc")) < tags.index(w("bidi")):
            raise ValueError(f"Invalid paragraph property order in {part_name}: w:bidi must precede w:jc")




def collect_template_placeholder_names(xml_bytes: bytes) -> list[str]:
    """أسماء المتغيرات المرئية فقط؛ لا تُقرأ تعليمات حقول Word."""
    root = etree.fromstring(xml_bytes)
    names: list[str] = []
    for para in root.iter(w("p")):
        if paragraph_has_nested_story(para):
            continue
        visible = "".join(node.text or "" for node in para.iter(w("t")))
        for match in VISIBLE_VARIABLE_RE.finditer(visible):
            name = visible_variable_name(match)
            if name and name not in names:
                names.append(name)
    return names




def repair_text_nodes(root: etree._Element) -> None:
    for text_node in root.iter(w("t")):
        if text_node.text:
            text_node.text = sanitize_xml_text(text_node.text, strip=False)
    for instr_node in root.iter(w("instrText")):
        if instr_node.text:
            instr_node.text = sanitize_xml_text(instr_node.text, strip=False)


def _rfonts_attr_keys(rfonts: etree._Element) -> list[tuple[str, str]]:
    """Return (attr_key, local_name) for font-related attributes on w:rFonts."""
    keys: list[tuple[str, str]] = []
    for key in rfonts.attrib:
        local = etree.QName(key).localname
        if local in ("ascii", "hAnsi", "eastAsia", "cs", "asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            keys.append((key, local))
    return keys


def _paragraph_own_text(para: etree._Element) -> str:
    """نص الفقرة نفسها فقط، دون فقرات متداخلة داخل رسومات/مربعات نص."""
    parts: list[str] = []
    for node in para.iter(w("t")):
        owning = node.getparent()
        while owning is not None and owning.tag != w("p"):
            owning = owning.getparent()
        if owning is para:
            parts.append(node.text or "")
    return "".join(parts)


def _paragraph_is_cover_footer(para: etree._Element) -> bool:
    text = _paragraph_own_text(para)
    return any(marker in text for marker in REPORT_COVER_FOOTER_MARKERS)


def _rewrite_rfonts_element(rfonts: etree._Element, font_name: str) -> bool:
    """Replace Cocon* font names with the target font. Leave other fonts untouched."""
    changed = False
    for key, local in list(_rfonts_attr_keys(rfonts)):
        value = rfonts.get(key) or ""
        if value in REPORT_COCON_FONT_NAMES:
            if local.endswith("Theme"):
                del rfonts.attrib[key]
                # ضع الخط الصريح بدل إشارة الثيم إن لزم.
                explicit = local.replace("Theme", "") if local.endswith("Theme") else local
                if explicit in ("ascii", "hAnsi", "eastAsia", "cs"):
                    rfonts.set(f"{{{W_NS}}}{explicit}", font_name)
                changed = True
            elif value != font_name:
                rfonts.set(key, font_name)
                changed = True
    return changed


def apply_report_fonts_to_part(xml_bytes: bytes) -> bytes:
    """
    اجعل خط التقرير Tajawal بدل Cocon، مع الإبقاء على CoconNextArabic-Light
    فقط في فقرات أسفل الغلاف (الرقم المرجعي + تاريخ التقرير).
    """
    root = etree.fromstring(xml_bytes)
    # أولاً: كل إشارات Cocon → Tajawal (بما فيها sdtPr / sdtEndPr وغيرها).
    for rfonts in root.iter(w("rFonts")):
        _rewrite_rfonts_element(rfonts, REPORT_BODY_FONT)
    # ثانياً: ثبّت Cocon على فقرات أسفل الغلاف فقط.
    for para in root.iter(w("p")):
        if not _paragraph_is_cover_footer(para):
            continue
        for rpr in para.iter(w("rPr")):
            parent = rpr.getparent()
            if parent is None or parent.tag not in (w("pPr"), w("r")):
                continue
            owning_para = parent
            while owning_para is not None and owning_para.tag != w("p"):
                owning_para = owning_para.getparent()
            if owning_para is not para:
                continue
            rfonts = rpr.find(w("rFonts"))
            if rfonts is None:
                rfonts = etree.SubElement(rpr, w("rFonts"))
            for key, local in list(_rfonts_attr_keys(rfonts)):
                if local.endswith("Theme"):
                    del rfonts.attrib[key]
            rfonts.set(f"{{{W_NS}}}cs", REPORT_COVER_FOOTER_FONT)
            rfonts.set(f"{{{W_NS}}}ascii", REPORT_COVER_FOOTER_FONT)
            rfonts.set(f"{{{W_NS}}}hAnsi", REPORT_COVER_FOOTER_FONT)
    return serialize_word_xml(root)


def ensure_tajawal_in_font_table(xml_bytes: bytes) -> bytes:
    """أضف Tajawal إلى جدول الخطوط إن لم يكن موجوداً."""
    root = etree.fromstring(xml_bytes)
    for font in root.iter(w("font")):
        name = font.get(f"{{{W_NS}}}name") or font.get("name") or ""
        if name == REPORT_BODY_FONT:
            return xml_bytes
    font_el = etree.SubElement(root, w("font"))
    font_el.set(f"{{{W_NS}}}name", REPORT_BODY_FONT)
    charset = etree.SubElement(font_el, w("charset"))
    charset.set(f"{{{W_NS}}}val", "00")
    family = etree.SubElement(font_el, w("family"))
    family.set(f"{{{W_NS}}}val", "swiss")
    pitch = etree.SubElement(font_el, w("pitch"))
    pitch.set(f"{{{W_NS}}}val", "variable")
    return serialize_word_xml(root)


def apply_tajawal_to_styles(xml_bytes: bytes) -> bytes:
    """حوّل إشارات Cocon في الأنماط إلى Tajawal."""
    root = etree.fromstring(xml_bytes)
    for rfonts in root.iter(w("rFonts")):
        _rewrite_rfonts_element(rfonts, REPORT_BODY_FONT)
    return serialize_word_xml(root)


def apply_visible_variable_values(
    xml_bytes: bytes,
    text_values: dict[str, str] | None = None,
) -> tuple[bytes, int, int]:
    root = etree.fromstring(xml_bytes)
    repair_text_nodes(root)
    clean_text_values = text_values or {}
    found, filled = replace_visible_variables(root, clean_text_values)
    flatten_mail_merge_fields(root)

    out = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    return out, found, filled


















def _save_print_jpeg(img: Image.Image, quality: int = DOCUMENT_IMAGE_JPEG_QUALITY) -> bytes:
    """JPEG أساسي (baseline) — python-docx يرفض غالباً Progressive/mozjpeg المعقّد."""
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    elif img.mode == "L":
        img = img.convert("RGB")
    out = io.BytesIO()
    img.save(
        out,
        format="JPEG",
        quality=max(60, min(100, int(quality))),
        optimize=False,
        progressive=False,
        subsampling=0,  # 4:4:4 — أوضح للنصوص والجداول
    )
    return out.getvalue()


def _save_docx_safe_png(img: Image.Image) -> bytes:
    if img.mode not in ("RGB", "RGBA", "L"):
        img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="PNG", optimize=False)
    return out.getvalue()


















def strip_mail_merge_settings(xml_bytes: bytes) -> bytes:
    """أزل اتصال Mail Merge القديم فقط، دون لمس إعدادات تحديث الفهرس أو الحقول."""
    root = etree.fromstring(xml_bytes)
    for mail_merge in list(root.iter(w("mailMerge"))):
        parent = mail_merge.getparent()
        if parent is not None:
            parent.remove(mail_merge)
    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )


def strip_mail_merge_relationships(xml_bytes: bytes) -> bytes:
    root = etree.fromstring(xml_bytes)
    for rel in list(root):
        rel_type = (rel.get("Type") or "").lower()
        target = (rel.get("Target") or "").lower()
        target_mode = (rel.get("TargetMode") or "").lower()
        is_merge_relation = (
            rel_type.endswith("/mailmergesource")
            or rel_type.endswith("/recipientdata")
            or "recipientdata" in target
            or (target_mode == "external" and ("xlsx" in target or "projects" in target))
        )
        if is_merge_relation:
            root.remove(rel)
    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )


def strip_recipient_content_type(xml_bytes: bytes) -> bytes:
    root = etree.fromstring(xml_bytes)
    for child in list(root):
        part_name = (child.get("PartName") or "").lower()
        if part_name == "/word/recipientdata.xml":
            root.remove(child)
    return etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )


def write_docx_zip(
    zin: zipfile.ZipFile,
    modified: dict[str, bytes],
    removed: set[str] | None = None,
) -> bytes:
    out_buf = io.BytesIO()
    removed_names = removed or set()
    names = zin.namelist()
    ordered: list[str] = []
    if "mimetype" in names and "mimetype" not in removed_names:
        ordered.append("mimetype")
    for name in names:
        if name != "mimetype" and name not in removed_names:
            ordered.append(name)
    for path in modified:
        if path not in ordered and path not in removed_names:
            ordered.append(path)

    with zipfile.ZipFile(out_buf, "w") as zout:
        for fname in ordered:
            if fname in modified:
                data = modified[fname]
            elif fname in names:
                data = zin.read(fname)
            else:
                continue
            compress = zipfile.ZIP_STORED if fname == "mimetype" else zipfile.ZIP_DEFLATED
            zi = zipfile.ZipInfo(fname)
            zi.compress_type = compress
            zout.writestr(zi, data)
    return out_buf.getvalue()


def downscale_jpeg_bytes(
    img_bytes: bytes,
    max_dimension: int,
    *,
    quality: int = DOCUMENT_IMAGE_JPEG_QUALITY,
    max_width: int | None = None,
    max_height: int | None = None,
    force_reencode: bool = False,
) -> bytes:
    """يحدّ الأبعاد ويعيد ترميز JPEG أساسي متوافق مع python-docx عند الحاجة."""
    try:
        img = Image.open(io.BytesIO(img_bytes))
        img.load()
        source_format = img.format
        img = img.convert("RGB")
        width, height = img.size
        if width <= 0 or height <= 0:
            raise ValueError(f"invalid image size {width}x{height}")
        scale = 1.0
        if max_dimension > 0:
            longest = max(width, height)
            if longest > max_dimension:
                scale = min(scale, max_dimension / longest)
        if max_width and width > max_width:
            scale = min(scale, max_width / width)
        if max_height and height > max_height:
            scale = min(scale, max_height / height)
        if scale < 1.0:
            img = img.resize(
                (max(1, int(width * scale)), max(1, int(height * scale))),
                Image.LANCZOS,
            )
            return _save_print_jpeg(img, quality)
        # حسابات القيمة: دائماً أعد الترميز إلى baseline JPEG حتى لو المصدر JPEG
        # (ملفات sharp/mozjpeg progressive كانت تسبب UnrecognizedImageError).
        if force_reencode or source_format != "JPEG":
            return _save_print_jpeg(img, quality)
        return img_bytes
    except Exception as exc:
        log_exception("downscale_jpeg_bytes failed", exc)
        return img_bytes


def prepare_valuation_image_bytes(img_bytes: bytes) -> tuple[bytes, int, int]:
    """
    إعداد صورة حسابات قيمة للطباعة:
    - دائماً JPEG أساسي عبر Pillow (متوافق مع python-docx)
    - تخطي ترميز sharp كان يسبب UnrecognizedImageError ثم PNG ضخم بطيء جداً
    يعيد (bytes, width_px, height_px)
    """
    if not img_bytes or len(img_bytes) < 32:
        raise ValueError(f"valuation image empty or too small ({0 if not img_bytes else len(img_bytes)} bytes)")

    try:
        img = Image.open(io.BytesIO(img_bytes))
        img.load()
        img = img.convert("RGB")
        width, height = img.size
        if width <= 0 or height <= 0:
            raise ValueError(f"invalid valuation size {width}x{height}")
        scale = 1.0
        longest = max(width, height)
        if longest > VALUATION_IMAGE_MAX_DIMENSION_PX:
            scale = min(scale, VALUATION_IMAGE_MAX_DIMENSION_PX / longest)
        if width > 4800:
            scale = min(scale, 4800 / width)
        if height > 14000:
            scale = min(scale, 14000 / height)
        if scale < 1.0:
            width = max(1, int(width * scale))
            height = max(1, int(height * scale))
            img = img.resize((width, height), Image.LANCZOS)
        prepared = _save_print_jpeg(img, VALUATION_IMAGE_JPEG_QUALITY)
        if not prepared or len(prepared) < 32:
            raise ValueError("valuation image re-encode produced empty output")
        return prepared, width, height
    except Exception:
        # مسار احتياطي بنفس الضمانات
        prepared = downscale_jpeg_bytes(
            img_bytes,
            VALUATION_IMAGE_MAX_DIMENSION_PX,
            quality=VALUATION_IMAGE_JPEG_QUALITY,
            max_width=4800,
            max_height=14000,
            force_reencode=True,
        )
        if not prepared or len(prepared) < 32:
            raise ValueError("valuation image re-encode produced empty output")
        with Image.open(io.BytesIO(prepared)) as check:
            check.load()
            return prepared, check.size[0], check.size[1]


def prepare_valuation_image_png_fallback(img_bytes: bytes) -> bytes:
    """احتياطي نادر جداً — يُفضَّل JPEG؛ يُستدعى فقط إن رُفض JPEG بعد Pillow."""
    img = Image.open(io.BytesIO(img_bytes))
    img.load()
    img = img.convert("RGB")
    return _save_docx_safe_png(img)


def crop_to_fill_jpeg_bytes(
    img_bytes: bytes,
    target_width: int,
    target_height: int,
) -> bytes:
    """إرث: قصّ لملء الإطار (cover). صور الأصول تستخدم ‎stretch_to_fill_canvas_jpeg_bytes‎."""
    img = Image.open(io.BytesIO(img_bytes))
    img = img.convert("RGB")
    width, height = img.size
    target_ratio = max(1, target_width) / max(1, target_height)
    source_ratio = width / max(1, height)
    if source_ratio > target_ratio:
        crop_width = max(1, int(height * target_ratio))
        left = max(0, (width - crop_width) // 2)
        cropped = img.crop((left, 0, left + crop_width, height))
    else:
        crop_height = max(1, int(width / target_ratio))
        top = max(0, (height - crop_height) // 2)
        cropped = img.crop((0, top, width, top + crop_height))
    longest = max(cropped.size)
    if longest > ASSET_IMAGE_MAX_SQUARE_PX:
        scale = ASSET_IMAGE_MAX_SQUARE_PX / longest
        cropped = cropped.resize(
            (max(1, int(cropped.width * scale)), max(1, int(cropped.height * scale))),
            Image.LANCZOS,
        )
    return _save_print_jpeg(cropped, ASSET_IMAGE_JPEG_QUALITY)


def _canvas_pixel_size_for_cell(
    target_width_emu: int,
    target_height_emu: int,
    max_side_px: int = ASSET_IMAGE_MAX_SQUARE_PX,
) -> tuple[int, int]:
    """أبعاد بكسل للوحة بنفس نسبة الخلية، مع حد أقصى للضلع الأطول."""
    tw = max(1, int(target_width_emu))
    th = max(1, int(target_height_emu))
    side = max(64, int(max_side_px))
    if tw >= th:
        canvas_w = side
        canvas_h = max(1, int(round(side * th / tw)))
    else:
        canvas_h = side
        canvas_w = max(1, int(round(side * tw / th)))
    return canvas_w, canvas_h


def _high_quality_stretch(img: Image.Image, canvas_w: int, canvas_h: int) -> Image.Image:
    """
    تمطيط (stretch) إلى مقاس ثابت — محورا العرض والارتفاع يُعدَّلان بشكل مستقل
    لملء الإطار بالكامل بدون قصّ وبدون فراغات.
    خطوة تصغير واحدة وسيطة كافية للصور الكبيرة جداً؛ تجنّب حلقات LANCZOS المتعددة
    التي كانت تبطّئ دمج مئات الصور وتستهلك الذاكرة على خوادم 4GB.
    """
    if img.size == (canvas_w, canvas_h):
        return img
    work = img
    if work.width > canvas_w * 3 and work.height > canvas_h * 3:
        work = work.resize(
            (max(canvas_w, work.width // 2), max(canvas_h, work.height // 2)),
            Image.BILINEAR,
        )
    if work.size != (canvas_w, canvas_h):
        work = work.resize((canvas_w, canvas_h), Image.LANCZOS)
    return work


def stretch_to_fill_canvas_jpeg_bytes(
    img_bytes: bytes,
    target_width_emu: int,
    target_height_emu: int,
    *,
    max_side_px: int = ASSET_IMAGE_MAX_SQUARE_PX,
    quality: int = ASSET_IMAGE_JPEG_QUALITY,
) -> bytes:
    """
    توحيد مساحة صور الأصول بدون اقتطاع وبدون هوامش داخلية:
    - الصورة كاملة تُمطَّط لتملأ الخلية 100%
    - كل الخلايا بنفس المقاس
    - المصدر مضغوط مسبقاً من Node؛ إعادة عيّنة واحدة عند الحاجة فقط
    """
    try:
        img = Image.open(io.BytesIO(img_bytes))
        img.load()
        if img.mode != "RGB":
            img = img.convert("RGB")
    except Exception:
        return img_bytes

    src_w, src_h = img.size
    if src_w <= 0 or src_h <= 0:
        return img_bytes

    canvas_w, canvas_h = _canvas_pixel_size_for_cell(
        target_width_emu,
        target_height_emu,
        max_side_px,
    )
    # إن طابقت الأبعاد اللوحة: ترميز Pillow آمن لـ python-docx فقط (بدون إعادة تحجيم)
    if src_w == canvas_w and src_h == canvas_h:
        return _save_print_jpeg(img, quality)
    stretched = _high_quality_stretch(img, canvas_w, canvas_h)
    return _save_print_jpeg(stretched, quality)


# توافق مع الاستدعاءات/الاختبارات القديمة إن وُجدت
def fit_contain_to_canvas_jpeg_bytes(
    img_bytes: bytes,
    target_width_emu: int,
    target_height_emu: int,
    *,
    max_side_px: int = ASSET_IMAGE_MAX_SQUARE_PX,
    quality: int = ASSET_IMAGE_JPEG_QUALITY,
    fill_rgb: tuple[int, int, int] = (255, 255, 255),
) -> bytes:
    del fill_rgb  # لم يعد مستخدماً — التمطيط يملأ الإطار
    return stretch_to_fill_canvas_jpeg_bytes(
        img_bytes,
        target_width_emu,
        target_height_emu,
        max_side_px=max_side_px,
        quality=quality,
    )


def package_base_for_rels(rels_name: str) -> str:
    if rels_name == "_rels/.rels":
        return ""
    marker = "/_rels/"
    if marker not in rels_name:
        return ""
    prefix, rel_file = rels_name.split(marker, 1)
    source_part = f"{prefix}/{rel_file[:-5]}" if rel_file.endswith(".rels") else f"{prefix}/{rel_file}"
    return source_part.rsplit("/", 1)[0] if "/" in source_part else ""


def validate_docx_package(docx_bytes: bytes) -> None:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        bad = zf.testzip()
        if bad:
            raise ValueError(f"Corrupt zip member: {bad}")
        names = set(zf.namelist())
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise ValueError("Invalid docx package: required parts are missing")

        ct_root = etree.fromstring(zf.read("[Content_Types].xml"))
        ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
        defaults: set[str] = set()
        overrides: set[str] = set()
        for child in ct_root:
            if child.tag == f"{{{ct_ns}}}Default":
                ext = (child.get("Extension") or "").lower()
                if ext in defaults:
                    raise ValueError(f"Duplicate content type default: {ext}")
                defaults.add(ext)
            elif child.tag == f"{{{ct_ns}}}Override":
                part = child.get("PartName") or ""
                if part in overrides:
                    raise ValueError(f"Duplicate content type override: {part}")
                overrides.add(part)

        for name in names:
            if name.endswith("/"):
                continue
            ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
            if f"/{name}" not in overrides and ext not in defaults:
                raise ValueError(f"Missing content type for {name}")
            if name.endswith(".xml") or name.endswith(".rels"):
                etree.fromstring(zf.read(name))

        rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
        for rels_name in [n for n in names if n.endswith(".rels")]:
            root = etree.fromstring(zf.read(rels_name))
            ids: set[str] = set()
            base = package_base_for_rels(rels_name)
            for rel in root:
                if rel.tag != f"{{{rel_ns}}}Relationship":
                    continue
                rid = rel.get("Id") or ""
                if rid in ids:
                    raise ValueError(f"Duplicate relationship id {rid} in {rels_name}")
                ids.add(rid)
                if (rel.get("TargetMode") or "").lower() == "external":
                    continue
                target = (rel.get("Target") or "").split("#", 1)[0]
                if not target:
                    continue
                part = target.lstrip("/") if target.startswith("/") else posixpath.normpath(posixpath.join(base, target))
                if part not in names:
                    raise ValueError(f"Missing relationship target {part} from {rels_name}")

        for part_name in [n for n in names if MERGE_PARTS_RE.match(n)]:
            validate_part_xml(zf.read(part_name), part_name)


def _renumber_drawing_ids_in_xml(raw: bytes, next_id: int) -> tuple[bytes | None, int]:
    try:
        root = etree.fromstring(raw)
    except Exception:
        return None, next_id
    changed = False
    for el in root.iter():
        local = etree.QName(el).localname
        if local not in ("docPr", "cNvPr") or el.get("id") is None:
            continue
        el.set("id", str(next_id))
        next_id += 1
        changed = True
    if not changed:
        return None, next_id
    return (
        etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"),
        next_id,
    )


def normalize_docx_drawing_ids(docx_bytes: bytes) -> bytes:
    """يعيد ترقيم معرّفات الرسم لتفادي التعارض مع صور الرأس/التذييل (مشكلة Word شائعة)."""
    modified: dict[str, bytes] = {}
    next_id = 1
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        for name in zin.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            rewritten, next_id = _renumber_drawing_ids_in_xml(zin.read(name), next_id)
            if rewritten is not None:
                modified[name] = rewritten
        if not modified:
            return docx_bytes
        return write_docx_zip(zin, modified)


def normalize_docx_drawing_ids_inplace(docx_path: str) -> None:
    """
    تطبيع معرّفات الرسم على القرص.
    الصور (media) تُحفظ ZIP_STORED لأنها JPEG مضغوطة أصلاً — أسرع بكثير من إعادة deflate لمئات الملفات.
    """
    tmp_path = f"{docx_path}.norm.tmp"
    next_id = 1
    changed_any = False
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w") as zout:
        for info in zin.infolist():
            name = info.filename
            is_media = name.startswith("word/media/")
            is_xml = name.startswith("word/") and name.endswith(".xml")
            data = zin.read(name)
            if is_xml:
                rewritten, next_id = _renumber_drawing_ids_in_xml(data, next_id)
                if rewritten is not None:
                    data = rewritten
                    changed_any = True
            if name == "mimetype" or is_media:
                zout.writestr(info, data, compress_type=zipfile.ZIP_STORED)
            else:
                zout.writestr(info, data, compress_type=zipfile.ZIP_DEFLATED)
    if changed_any:
        os.replace(tmp_path, docx_path)
    else:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def docx_qn(tag: str) -> str:
    from docx.oxml.ns import qn

    return qn(tag)


def make_docx_element(tag: str):
    from docx.oxml import OxmlElement

    return OxmlElement(tag)


def set_docx_cell_margins(cell, margin_dxa: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(docx_qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = make_docx_element("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        node = make_docx_element(f"w:{side}")
        node.set(docx_qn("w:w"), str(margin_dxa))
        node.set(docx_qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_docx_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(docx_qn("w:tblBorders"))
    if tbl_borders is not None:
        tbl_pr.remove(tbl_borders)
    tbl_borders = make_docx_element("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = make_docx_element(f"w:{border_name}")
        border.set(docx_qn("w:val"), "none")
        border.set(docx_qn("w:sz"), "0")
        border.set(docx_qn("w:space"), "0")
        border.set(docx_qn("w:color"), "auto")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def set_docx_table_cell_spacing(table, spacing_dxa: int = 0) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(docx_qn("w:tblCellSpacing"))
    if existing is not None:
        tbl_pr.remove(existing)
    spacing = make_docx_element("w:tblCellSpacing")
    spacing.set(docx_qn("w:w"), str(max(0, int(spacing_dxa))))
    spacing.set(docx_qn("w:type"), "dxa")
    tbl_pr.append(spacing)


def set_docx_table_width(table, width_emu: int) -> None:
    width_dxa = max(1, int(width_emu / EMU_PER_INCH * 1440))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(docx_qn("w:tblW"))
    if tbl_w is not None:
        tbl_pr.remove(tbl_w)
    tbl_w = make_docx_element("w:tblW")
    tbl_w.set(docx_qn("w:w"), str(width_dxa))
    tbl_w.set(docx_qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    layout = make_docx_element("w:tblLayout")
    layout.set(docx_qn("w:type"), "fixed")
    tbl_pr.append(layout)


def set_docx_table_indent(table, indent_emu: int) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(docx_qn("w:tblInd"))
    if existing is not None:
        tbl_pr.remove(existing)
    indent = make_docx_element("w:tblInd")
    indent.set(docx_qn("w:w"), str(emu_to_twips(indent_emu)))
    indent.set(docx_qn("w:type"), "dxa")
    tbl_pr.append(indent)


def detach_docx_body_element(elem) -> Any:
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)
    return elem


def make_docx_title_element(doc, title: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(title)
    return detach_docx_body_element(p._element)


def make_docx_page_break_element(doc):
    from docx.enum.text import WD_BREAK

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return detach_docx_body_element(p._element)


def document_page_inner_box_emu(doc) -> tuple[int, int]:
    section = doc.sections[0]
    page_width = int(section.page_width or int(8.27 * EMU_PER_INCH))
    page_height = int(section.page_height or int(11.69 * EMU_PER_INCH))
    left_margin = int(section.left_margin or int(0.5 * EMU_PER_INCH))
    right_margin = int(section.right_margin or int(0.5 * EMU_PER_INCH))
    top_margin = int(section.top_margin or int(0.5 * EMU_PER_INCH))
    bottom_margin = int(section.bottom_margin or int(0.5 * EMU_PER_INCH))
    content_width = int(page_width - left_margin - right_margin)
    content_height = int(page_height - top_margin - bottom_margin)
    if content_width <= 0:
        content_width = int(7.28 * EMU_PER_INCH)
    if content_height <= 0:
        content_height = int(10.0 * EMU_PER_INCH)
    return content_width, content_height


def document_physical_page_box_emu(doc) -> tuple[int, int]:
    section = doc.sections[0]
    page_width = int(section.page_width or int(8.27 * EMU_PER_INCH))
    page_height = int(section.page_height or int(11.69 * EMU_PER_INCH))
    if page_width <= 0:
        page_width = int(8.27 * EMU_PER_INCH)
    if page_height <= 0:
        page_height = int(11.69 * EMU_PER_INCH)
    return page_width, page_height


def document_section_margins_emu(doc) -> tuple[int, int, int, int]:
    section = doc.sections[0]
    left_margin = int(section.left_margin or int(0.5 * EMU_PER_INCH))
    right_margin = int(section.right_margin or int(0.5 * EMU_PER_INCH))
    top_margin = int(section.top_margin or int(0.5 * EMU_PER_INCH))
    bottom_margin = int(section.bottom_margin or int(0.5 * EMU_PER_INCH))
    return left_margin, right_margin, top_margin, bottom_margin


def w_attr_int(el, name: str, fallback: int) -> int:
    if el is None:
        return fallback
    raw = el.get(docx_qn(f"w:{name}"))
    try:
        return int(raw) if raw is not None else fallback
    except (TypeError, ValueError):
        return fallback


def twips_to_emu(value: int) -> int:
    return int(value * EMU_PER_INCH / 1440)


def emu_to_twips(value: int | float) -> int:
    return int(round(float(value) * 1440 / EMU_PER_INCH))


def section_metrics_from_sect_pr(sect_pr) -> tuple[int, int, int, int, int, int] | None:
    if sect_pr is None:
        return None

    pg_sz = sect_pr.find(docx_qn("w:pgSz"))
    pg_mar = sect_pr.find(docx_qn("w:pgMar"))
    page_width_twips = w_attr_int(pg_sz, "w", 0)
    page_height_twips = w_attr_int(pg_sz, "h", 0)

    page_width = twips_to_emu(page_width_twips) if page_width_twips > 0 else DEFAULT_PAGE_WIDTH_EMU
    page_height = twips_to_emu(page_height_twips) if page_height_twips > 0 else DEFAULT_PAGE_HEIGHT_EMU
    left_margin = twips_to_emu(w_attr_int(pg_mar, "left", emu_to_twips(DEFAULT_PAGE_MARGIN_EMU)))
    right_margin = twips_to_emu(w_attr_int(pg_mar, "right", emu_to_twips(DEFAULT_PAGE_MARGIN_EMU)))
    top_margin = twips_to_emu(w_attr_int(pg_mar, "top", emu_to_twips(DEFAULT_PAGE_MARGIN_EMU)))
    bottom_margin = twips_to_emu(w_attr_int(pg_mar, "bottom", emu_to_twips(DEFAULT_PAGE_MARGIN_EMU)))
    return page_width, page_height, left_margin, right_margin, top_margin, bottom_margin


def first_section_metrics_at_or_after(children: list[Any], start_idx: int):
    for child in children[max(0, start_idx) :]:
        sect_pr = child if getattr(child, "tag", None) == docx_qn("w:sectPr") else child.find(f".//{docx_qn('w:sectPr')}")
        metrics = section_metrics_from_sect_pr(sect_pr)
        if metrics is not None:
            return metrics
    return None


def portrait_a4_metrics() -> tuple[int, int, int, int, int, int]:
    """مقاييس صفحة A4 طولية (مثل القالب الأساسي mv-word-template)."""
    return (
        DEFAULT_PAGE_WIDTH_EMU,
        DEFAULT_PAGE_HEIGHT_EMU,
        DEFAULT_PAGE_MARGIN_EMU,
        DEFAULT_PAGE_MARGIN_EMU,
        DEFAULT_PAGE_MARGIN_EMU,
        DEFAULT_PAGE_MARGIN_EMU,
    )


def is_landscape_metrics(metrics: tuple[int, int, int, int, int, int] | None) -> bool:
    if metrics is None:
        return False
    return metrics[0] > metrics[1]


def ensure_portrait_metrics(
    metrics: tuple[int, int, int, int, int, int] | None,
) -> tuple[int, int, int, int, int, int]:
    """أبقِ الهوامش إن أمكن لكن افرض عرضاً أصغر من الارتفاع (طولي)."""
    if metrics is None:
        return portrait_a4_metrics()
    page_width, page_height, left_m, right_m, top_m, bottom_m = metrics
    if page_width > page_height:
        page_width, page_height = page_height, page_width
    if page_width <= 0 or page_height <= 0:
        return portrait_a4_metrics()
    return page_width, page_height, left_m, right_m, top_m, bottom_m


def apply_portrait_a4_to_sect_pr(sect_pr) -> None:
    """حوّل إعدادات القسم إلى A4 طولي بدون اتجاه landscape."""
    if sect_pr is None:
        return
    pg_sz = sect_pr.find(docx_qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = make_docx_element("w:pgSz")
        sect_pr.insert(0, pg_sz)
    pg_sz.set(docx_qn("w:w"), str(emu_to_twips(DEFAULT_PAGE_WIDTH_EMU)))
    pg_sz.set(docx_qn("w:h"), str(emu_to_twips(DEFAULT_PAGE_HEIGHT_EMU)))
    orient_key = docx_qn("w:orient")
    if orient_key in pg_sz.attrib:
        del pg_sz.attrib[orient_key]

    pg_mar = sect_pr.find(docx_qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = make_docx_element("w:pgMar")
        sect_pr.append(pg_mar)
    margin_twips = str(emu_to_twips(DEFAULT_PAGE_MARGIN_EMU))
    for side in ("top", "right", "bottom", "left", "header", "footer", "gutter"):
        if pg_mar.get(docx_qn(f"w:{side}")) is None:
            pg_mar.set(docx_qn(f"w:{side}"), margin_twips)


def body_final_sect_pr(body):
    children = list(body)
    if not children:
        return None
    last = children[-1]
    if getattr(last, "tag", None) == docx_qn("w:sectPr"):
        return last
    return None


def make_section_break_paragraph_from_sect_pr(sect_pr, *, next_page: bool = True):
    """فقرة فاصل أقسام: الـ sectPr هنا يغلق القسم السابق بخصائصه."""
    break_p = make_docx_element("w:p")
    p_pr = make_docx_element("w:pPr")
    sect_copy = deepcopy(sect_pr)
    if next_page:
        type_el = sect_copy.find(docx_qn("w:type"))
        if type_el is None:
            type_el = make_docx_element("w:type")
            sect_copy.insert(0, type_el)
        type_el.set(docx_qn("w:val"), "nextPage")
    p_pr.append(sect_copy)
    break_p.append(p_pr)
    return break_p




def document_valuation_image_layout_emu(
    doc,
    section_metrics: tuple[int, int, int, int, int, int] | None = None,
) -> tuple[int, int, int]:
    """Use 95% of the physical page width with equal physical side margins."""
    if section_metrics is None:
        page_width, _ = document_physical_page_box_emu(doc)
        left_margin, right_margin, _top_margin, _bottom_margin = document_section_margins_emu(doc)
    else:
        page_width, _page_height, left_margin, right_margin, _top_margin, _bottom_margin = section_metrics
    page_width = max(1, page_width)
    target_width = max(1, int(page_width * IMAGE_CONTENT_WIDTH_RATIO))
    physical_side_margin = max(0, (page_width - target_width) // 2)
    return (
        target_width,
        physical_side_margin - left_margin,
        physical_side_margin - right_margin,
    )


def configure_rtl_valuation_image_paragraph(paragraph, indent_left: int, indent_right: int) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    for tag in ("w:jc", "w:bidi", "w:ind"):
        el = p_pr.find(docx_qn(tag))
        if el is not None:
            p_pr.remove(el)

    p_pr.append(make_docx_element("w:bidi"))
    jc = make_docx_element("w:jc")
    jc.set(docx_qn("w:val"), "right")
    p_pr.append(jc)

    ind = make_docx_element("w:ind")
    left_twips = emu_to_twips(indent_left)
    right_twips = emu_to_twips(indent_right)
    ind.set(docx_qn("w:left"), str(left_twips))
    ind.set(docx_qn("w:right"), str(right_twips))
    ind.set(docx_qn("w:start"), str(right_twips))
    ind.set(docx_qn("w:end"), str(left_twips))
    p_pr.append(ind)


def scaled_image_size_for_width_only(
    img_bytes: bytes,
    target_width_emu: int,
    max_height_emu: int | None = None,
    *,
    width_px: int | None = None,
    height_px: int | None = None,
) -> tuple[int, int]:
    if width_px is None or height_px is None:
        with Image.open(io.BytesIO(img_bytes)) as img:
            width_px, height_px = img.size
    if width_px <= 0 or height_px <= 0:
        raise ValueError("invalid image dimensions")
    cx = max(1, int(target_width_emu))
    cy = max(1, int(cx * height_px / width_px))
    limit = max_height_emu if max_height_emu and max_height_emu > 0 else MAX_DRAWING_HEIGHT_EMU
    if cy > limit:
        scale = limit / cy
        cx = max(1, int(cx * scale))
        cy = max(1, int(cy * scale))
    return cx, cy


def document_content_box_emu(
    doc,
    *,
    title_reserve_emu: int | None = None,
) -> tuple[int, int]:
    content_width, content_height = document_page_inner_box_emu(doc)
    reserve = (
        int(title_reserve_emu)
        if title_reserve_emu is not None
        else int(0.65 * EMU_PER_INCH)
    )
    content_height = int(content_height - max(0, reserve))
    if content_height <= 0:
        content_height = int(9.0 * EMU_PER_INCH)
    return content_width, content_height


def document_cell_dimensions_emu(
    doc,
    images_per_row: int,
    image_rows_per_page: int,
    section_metrics: tuple[int, int, int, int, int, int] | None = None,
    *,
    fill_content_height: bool = False,
    content_height_ratio: float = 1.0,
    bottom_margin_emu: int = 0,
    content_width_ratio: float = IMAGE_CONTENT_WIDTH_RATIO,
    gap_emu: int = ASSET_IMAGE_GAP_EMU,
    title_reserve_emu: int | None = None,
) -> tuple[int, int]:
    _content_width, content_height = document_content_box_emu(
        doc,
        title_reserve_emu=title_reserve_emu,
    )
    if fill_content_height:
        usable_height = max(
            1,
            int(content_height * content_height_ratio) - max(0, int(bottom_margin_emu)),
        )
        content_height = usable_height
    if section_metrics is None:
        page_width, _page_height = document_physical_page_box_emu(doc)
    else:
        page_width = section_metrics[0]
    width_ratio = min(1.0, max(0.5, float(content_width_ratio)))
    gap = max(0, int(gap_emu))
    available_width = max(1, int(page_width * width_ratio))
    width_fit = max(1, (available_width - (images_per_row - 1) * gap) // images_per_row)
    height_fit = max(1, (content_height - (image_rows_per_page - 1) * gap) // image_rows_per_page)
    if fill_content_height:
        # مستندات العميل: املأ الارتفاع المتاح دون إجبار الخلية على مربع.
        return width_fit, height_fit
    return width_fit, max(1, min(width_fit, height_fit))


def scaled_image_size_for_width(img_bytes: bytes, target_width_emu: int, max_height_emu: int) -> tuple[int, int]:
    with Image.open(io.BytesIO(img_bytes)) as img:
        width_px, height_px = img.size
    if width_px <= 0 or height_px <= 0:
        raise ValueError("invalid image dimensions")
    cx = max(1, int(target_width_emu))
    cy = max(1, int(cx * height_px / width_px))
    if cy > max_height_emu > 0:
        scale = max_height_emu / cy
        cx = max(1, int(cx * scale))
        cy = max(1, int(cy * scale))
    return cx, cy


def make_docx_image_table_element(
    doc,
    images: list[ImageSource],
    images_per_row: int,
    image_rows_per_page: int,
    section_metrics: tuple[int, int, int, int, int, int] | None = None,
    *,
    fill_content_height: bool = False,
    content_height_ratio: float = 1.0,
    bottom_margin_emu: int = 0,
    contain_images: bool = False,
    content_width_ratio: float = IMAGE_CONTENT_WIDTH_RATIO,
    gap_dxa: int = ASSET_IMAGE_GAP_DXA,
    gap_emu: int = ASSET_IMAGE_GAP_EMU,
    title_reserve_emu: int | None = None,
    asset_max_side_px: int = ASSET_IMAGE_MAX_SQUARE_PX,
) -> tuple[Any, int]:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu, Pt

    rows = max(1, math.ceil(len(images) / images_per_row))
    cell_width_emu, image_emu = document_cell_dimensions_emu(
        doc,
        images_per_row,
        image_rows_per_page,
        section_metrics,
        fill_content_height=fill_content_height,
        content_height_ratio=content_height_ratio,
        bottom_margin_emu=bottom_margin_emu,
        content_width_ratio=content_width_ratio,
        gap_emu=gap_emu,
        title_reserve_emu=title_reserve_emu,
    )
    cell_width_dxa = max(1, int(cell_width_emu / EMU_PER_INCH * 1440))
    width_ratio = min(1.0, max(0.5, float(content_width_ratio)))
    gap_dxa_val = max(0, int(gap_dxa))
    gap_emu_val = max(0, int(gap_emu))

    table = doc.add_table(rows=rows, cols=images_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_docx_table_borders_none(table)
    set_docx_table_cell_spacing(table, gap_dxa_val)
    set_docx_table_width(
        table,
        cell_width_emu * images_per_row + (images_per_row - 1) * gap_emu_val,
    )
    if section_metrics is None:
        page_width, _page_height = document_physical_page_box_emu(doc)
        left_margin, _right_margin, _top_margin, _bottom_margin = document_section_margins_emu(doc)
    else:
        page_width, _page_height, left_margin, _right_margin, _top_margin, _bottom_margin = section_metrics
    physical_side_margin = max(0, int(page_width * (1 - width_ratio) / 2))
    set_docx_table_indent(table, physical_side_margin - left_margin)

    inserted = 0
    img_index = 0
    for row in table.rows:
        row.height = Emu(image_emu)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = Emu(cell_width_emu)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(docx_qn("w:tcW"))
            if tc_w is not None:
                tc_pr.remove(tc_w)
            tc_w = make_docx_element("w:tcW")
            tc_w.set(docx_qn("w:w"), str(cell_width_dxa))
            tc_w.set(docx_qn("w:type"), "dxa")
            tc_pr.append(tc_w)
            set_docx_cell_margins(cell, 0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            if img_index >= len(images):
                continue
            img_bytes: bytes | None = None
            try:
                # المصدر جاهز JPEG من Node — لا تعيد الترميز مرتين قبل الإدراج
                img_bytes = resolve_image_bytes(images[img_index])
                if contain_images:
                    # مستندات العميل: ترميز Pillow واحد متوافق مع python-docx
                    img_bytes = ensure_jpeg_bytes(img_bytes, DOCUMENT_IMAGE_JPEG_QUALITY)
                    fit_w, fit_h = scaled_image_size_for_width(
                        img_bytes,
                        cell_width_emu,
                        min(image_emu, MAX_DRAWING_HEIGHT_EMU),
                    )
                    if fit_w <= 0 or fit_h <= 0:
                        raise ValueError(f"invalid fit size {fit_w}x{fit_h}")
                    para.add_run().add_picture(
                        io.BytesIO(img_bytes),
                        width=Emu(fit_w),
                        height=Emu(fit_h),
                    )
                else:
                    # صور الأصول: تمطيط لملء الخلية (يعيد الترميز مرة واحدة فقط)
                    stretched = stretch_to_fill_canvas_jpeg_bytes(
                        img_bytes,
                        cell_width_emu,
                        image_emu,
                        max_side_px=asset_max_side_px,
                        quality=ASSET_IMAGE_JPEG_QUALITY,
                    )
                    para.add_run().add_picture(
                        io.BytesIO(stretched),
                        width=Emu(max(1, cell_width_emu)),
                        height=Emu(max(1, min(image_emu, MAX_DRAWING_HEIGHT_EMU))),
                    )
                inserted += 1
            except Exception as exc:
                log_exception("image insert skipped", exc)
            finally:
                img_bytes = None
            img_index += 1

    # python-docx leaves the table attached to the document body; detach before insertion.
    return detach_docx_body_element(table._tbl), inserted


def make_docx_valuation_image_element(
    doc,
    image: ImageSource,
    section_metrics: tuple[int, int, int, int, int, int] | None = None,
    *,
    layout_cache: dict[str, Any] | None = None,
) -> tuple[Any, int]:
    from docx.shared import Emu, Pt

    p = doc.add_paragraph()
    raw = resolve_image_bytes(image)
    try:
        # بكسل عالي + JPEG Pillow أساسي (مرة واحدة) — بدون مسار PNG الثقيل
        img_bytes, width_px, height_px = prepare_valuation_image_bytes(raw)
        raw = None  # حرر مبكراً
        cache = layout_cache if layout_cache is not None else {}
        if "target_width" not in cache:
            target_width, indent_left, indent_right = document_valuation_image_layout_emu(
                doc, section_metrics
            )
            _content_width, content_height = document_content_box_emu(doc, title_reserve_emu=0)
            cache["target_width"] = target_width
            cache["indent_left"] = indent_left
            cache["indent_right"] = indent_right
            cache["max_height"] = min(MAX_DRAWING_HEIGHT_EMU, max(1, int(content_height * 0.98)))
        target_width = cache["target_width"]
        indent_left = cache["indent_left"]
        indent_right = cache["indent_right"]
        max_height = cache["max_height"]
        cx, cy = scaled_image_size_for_width_only(
            img_bytes,
            target_width,
            max_height,
            width_px=width_px,
            height_px=height_px,
        )
        configure_rtl_valuation_image_paragraph(p, indent_left, indent_right)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        try:
            p.add_run().add_picture(io.BytesIO(img_bytes), width=Emu(cx), height=Emu(cy))
        except Exception as jpeg_exc:
            # نادر بعد Pillow — PNG من البايتات الجاهزة (بدون إعادة تحجيم من المصدر)
            log_exception("valuation JPEG insert failed, trying PNG", jpeg_exc)
            for run in list(p.runs):
                el = run._element
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
            png_bytes = prepare_valuation_image_png_fallback(img_bytes)
            p.add_run().add_picture(io.BytesIO(png_bytes), width=Emu(cx), height=Emu(cy))
        return detach_docx_body_element(p._element), 1
    except Exception as exc:
        log_exception("Skipped valuation image", exc)
        return detach_docx_body_element(p._element), 0




def block_text(block) -> str:
    return sanitize_xml_text(
        "".join(node.text or "" for node in block.iter(docx_qn("w:t"))),
    )


def find_body_heading_index(
    children: list[Any],
    headings: tuple[str, ...] | list[str],
) -> tuple[int | None, str | None]:
    """ابحث من النهاية وفي فقرات body المباشرة فقط كي لا نلتقط عنوان الفهرس."""
    wanted = {
        normalize_heading_text(heading): heading
        for heading in headings
    }
    for idx in range(len(children) - 1, -1, -1):
        child = children[idx]
        if etree.QName(child).localname != "p":
            continue
        normalized = normalize_heading_text(block_text(child))
        for key, heading in wanted.items():
            if normalized == key or normalized.startswith(key):
                return idx, heading
    return None, None


def block_is_empty_insertion_spacer(block) -> bool:
    if etree.QName(block).localname != "p":
        return False
    if block_text(block):
        return False
    protected_tags = (
        docx_qn("w:br"),
        docx_qn("w:sectPr"),
        docx_qn("w:drawing"),
        docx_qn("w:pict"),
    )
    return not any(any(True for _ in block.iter(tag)) for tag in protected_tags)


def remove_empty_spacers_after(body, heading_idx: int) -> int:
    insert_at = heading_idx + 1
    while True:
        children = list(body)
        if insert_at >= len(children) or not block_is_empty_insertion_spacer(children[insert_at]):
            return insert_at
        body.remove(children[insert_at])


def block_is_placeholder_image(block) -> bool:
    if etree.QName(block).localname != "p":
        return False
    has_image = any(True for _ in block.iter(docx_qn("w:drawing"))) or any(True for _ in block.iter(docx_qn("w:pict")))
    if not has_image:
        return False
    text = "".join(node.text or "" for node in block.iter(docx_qn("w:t")))
    return not text.strip()


def insert_image_grid_pages(
    doc,
    body,
    insert_at: int,
    images: list[ImageSource],
    layout: str,
    images_per_row: int,
    images_per_page: int,
    section_metrics: tuple[int, int, int, int, int, int] | None,
    asset_max_side_px: int = ASSET_IMAGE_MAX_SQUARE_PX,
) -> tuple[int, int]:
    """يدرج صفحات صور عند insert_at. يعيد (عدد المدرج, موضع الإدراج التالي)."""
    inserted = 0

    def insert_elem(elem: Any) -> None:
        nonlocal insert_at
        body.insert(insert_at, elem)
        insert_at += 1

    if layout == "valuation_pages":
        layout_cache: dict[str, Any] = {}
        for image in images:
            image_elem, count = make_docx_valuation_image_element(
                doc, image, section_metrics, layout_cache=layout_cache
            )
            if count <= 0:
                continue
            if inserted > 0:
                insert_elem(make_docx_page_break_element(doc))
            insert_elem(image_elem)
            inserted += count
            if inserted % 40 == 0:
                gc.collect()
        return inserted, insert_at

    image_rows_per_page = max(1, math.ceil(images_per_page / images_per_row))
    fill_client = layout == "client_grid"
    page_number = 0
    for page_idx in range(0, len(images), images_per_page):
        if page_idx > 0:
            insert_elem(make_docx_page_break_element(doc))
        table_elem, count = make_docx_image_table_element(
            doc,
            images[page_idx : page_idx + images_per_page],
            images_per_row,
            image_rows_per_page,
            section_metrics,
            fill_content_height=fill_client,
            content_height_ratio=CLIENT_DOCS_CONTENT_HEIGHT_RATIO if fill_client else 1.0,
            bottom_margin_emu=CLIENT_DOCS_BOTTOM_MARGIN_EMU if fill_client else 0,
            contain_images=fill_client,
            content_width_ratio=CLIENT_DOCS_CONTENT_WIDTH_RATIO if fill_client else IMAGE_CONTENT_WIDTH_RATIO,
            gap_dxa=CLIENT_DOCS_GAP_DXA if fill_client else ASSET_IMAGE_GAP_DXA,
            gap_emu=CLIENT_DOCS_GAP_EMU if fill_client else ASSET_IMAGE_GAP_EMU,
            title_reserve_emu=CLIENT_DOCS_TITLE_RESERVE_EMU if fill_client else None,
            asset_max_side_px=asset_max_side_px,
        )
        insert_elem(table_elem)
        inserted += count
        page_number += 1
        if page_number % 10 == 0:
            gc.collect()
    return inserted, insert_at


def insert_images_after_section_heading(
    doc,
    field: str,
    images: list[ImageSource],
    layout: str = "asset_grid",
    images_per_row: int = IMAGES_PER_ROW,
    images_per_page: int = IMAGES_PER_PAGE,
    asset_max_side_px: int = ASSET_IMAGE_MAX_SQUARE_PX,
) -> int:
    body = doc.element.body
    children = list(body)
    if not images:
        log(f"image section {field!r}: no images to insert")
        return 0
    target_idx, found_heading = find_body_heading_index(
        children,
        IMAGE_SECTION_HEADINGS.get(field, ()),
    )
    if target_idx is None:
        log(
            f"image section heading for {field!r} "
            "was not found in document body"
        )
        return 0
    log(f"image section heading found for {field!r}: {found_heading!r}")
    section_metrics = first_section_metrics_at_or_after(children, target_idx)
    insert_at = remove_empty_spacers_after(body, target_idx)
    inserted, _ = insert_image_grid_pages(
        doc,
        body,
        insert_at,
        images,
        layout,
        images_per_row,
        images_per_page,
        section_metrics,
        asset_max_side_px,
    )
    return inserted


def insert_client_images_after_section_heading(
    doc,
    images: list[ImageSource],
    images_per_row: int = CLIENT_DOCS_IMAGES_PER_ROW,
    images_per_page: int = CLIENT_DOCS_IMAGES_PER_PAGE,
) -> int:
    """
    إدراج مستندات العميل بعد عنوان مرفق 3 فقط.

    لا تُستخدم أي علامات مرجعية ولا يُنشأ مرفق بديل عند غياب العنوان.
    """
    if not images:
        log("client image section: no images to insert")
        return 0

    body = doc.element.body
    children = list(body)
    target_idx, found_name = find_body_heading_index(
        children,
        IMAGE_SECTION_HEADINGS["client"],
    )
    if target_idx is None:
        log(
            "client image section heading was not found in document body"
        )
        return 0

    log(f"client image section heading found: {found_name!r}")
    raw_metrics = first_section_metrics_at_or_after(children, target_idx)
    section_metrics = ensure_portrait_metrics(raw_metrics)
    insert_at = remove_empty_spacers_after(body, target_idx)
    inserted, _ = insert_image_grid_pages(
        doc,
        body,
        insert_at,
        images,
        "client_grid",
        images_per_row,
        images_per_page,
        section_metrics,
    )
    return inserted


def apply_image_sections_docx_api(
    docx_bytes: bytes,
    asset_images: list[ImageSource],
    valuation_images: list[ImageSource],
    client_images: list[ImageSource] | None = None,
    images_per_row: int = IMAGES_PER_ROW,
    images_per_page: int = IMAGES_PER_PAGE,
    client_images_per_row: int = CLIENT_DOCS_IMAGES_PER_ROW,
    client_images_per_page: int = CLIENT_DOCS_IMAGES_PER_PAGE,
    output_path: str | None = None,
) -> tuple[bytes | None, dict[str, int]]:
    from docx import Document

    asset_max_side = adaptive_asset_max_side(len(asset_images))
    doc = Document(io.BytesIO(docx_bytes))
    stats = {"asset": 0, "valuation": 0, "client": 0}
    # الأهم أولاً: حسابات القيمة + مستندات العميل قبل مئات صور الأصول،
    # حتى لا تفشل إضافتها بعد تضخّم المستند/تعارض معرّفات الرسم.
    stats["valuation"] = insert_images_after_section_heading(
        doc,
        "valuation",
        valuation_images,
        "valuation_pages",
    )
    valuation_images.clear()
    gc.collect()
    log(f"valuation images inserted: {stats['valuation']}")

    stats["client"] = insert_client_images_after_section_heading(
        doc,
        client_images or [],
        client_images_per_row,
        client_images_per_page,
    )
    if client_images is not None:
        client_images.clear()
    gc.collect()
    log(f"client images inserted: {stats['client']}")

    stats["asset"] = insert_images_after_section_heading(
        doc,
        "asset",
        asset_images,
        "asset_grid",
        images_per_row,
        images_per_page,
        asset_max_side,
    )
    asset_images.clear()
    gc.collect()
    log(f"asset images inserted: {stats['asset']}")

    image_total = (
        (stats.get("asset") or 0)
        + (stats.get("valuation") or 0)
        + (stats.get("client") or 0)
    )

    if output_path:
        doc.save(output_path)
        del doc
        gc.collect()
        try:
            normalize_docx_drawing_ids_inplace(output_path)
        except Exception as exc:
            log_exception("normalize_docx_drawing_ids_inplace failed", exc)
        if image_total <= 40:
            try:
                with open(output_path, "rb") as fh:
                    validate_docx_package(fh.read())
            except Exception as exc:
                log_exception("validate_docx_package failed", exc)
        return None, stats

    out = io.BytesIO()
    doc.save(out)
    del doc
    result = normalize_docx_drawing_ids(out.getvalue())
    if image_total <= 40:
        validate_docx_package(result)
    else:
        with zipfile.ZipFile(io.BytesIO(result), "r") as zf:
            bad = zf.testzip()
            if bad:
                raise ValueError(f"Corrupt zip member: {bad}")
            names = set(zf.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ValueError("Invalid docx package: required parts are missing")
    return result, stats


def merge_package(payload: dict[str, Any]) -> bytes | None:
    global ASSET_IMAGE_JPEG_QUALITY
    global VALUATION_IMAGE_JPEG_QUALITY
    global DOCUMENT_IMAGE_JPEG_QUALITY

    output_path = str(payload.get("outputPath") or "").strip() or None
    template_path = str(payload.get("templatePath") or "").strip()
    template_b64 = payload.get("templateBase64") or ""

    if template_path and os.path.isfile(template_path):
        with open(template_path, "rb") as fh:
            template_bytes = fh.read()
    elif template_b64:
        template_bytes = base64.b64decode(template_b64)
    else:
        raise ValueError("templatePath or templateBase64 missing")

    text_values = payload.get("textValues") or {}
    report_preparers_present, report_preparers = collect_report_preparers(payload)
    asset_images = collect_image_sources(payload, "assetImagePaths", "assetImagesBase64")
    valuation_images = collect_image_sources(payload, "valuationImagePaths", "valuationImagesBase64")
    client_images = collect_image_sources(payload, "clientImagePaths", "clientImagesBase64")
    image_layout = payload.get("imageLayout") if isinstance(payload.get("imageLayout"), dict) else {}
    try:
        image_quality = max(60, min(100, int(image_layout.get("imageQuality", 90))))
    except (TypeError, ValueError):
        image_quality = 90
    ASSET_IMAGE_JPEG_QUALITY = adaptive_asset_quality(
        len(asset_images),
        image_quality,
    )
    VALUATION_IMAGE_JPEG_QUALITY = image_quality
    DOCUMENT_IMAGE_JPEG_QUALITY = image_quality
    try:
        images_per_row = max(1, min(6, int(image_layout.get("imagesPerRow", IMAGES_PER_ROW))))
    except (TypeError, ValueError):
        images_per_row = IMAGES_PER_ROW
    if images_per_row <= 1:
        auto_images_per_page = 2
    elif images_per_row == 2:
        auto_images_per_page = 4
    else:
        auto_images_per_page = images_per_row * (5 if images_per_row >= 4 else 4)
    # لصف واحد أو اثنين نفرض التخطيط المطلوب حتى لو وصلت قيمة قديمة من الواجهة
    if images_per_row <= 2:
        images_per_page = auto_images_per_page
    else:
        try:
            images_per_page = max(1, int(image_layout.get("imagesPerPage", auto_images_per_page)))
        except (TypeError, ValueError):
            images_per_page = auto_images_per_page
        images_per_page = max(images_per_row, images_per_page)
    try:
        client_images_per_row = int(image_layout.get("clientImagesPerRow", CLIENT_DOCS_IMAGES_PER_ROW))
    except (TypeError, ValueError):
        client_images_per_row = CLIENT_DOCS_IMAGES_PER_ROW
    if client_images_per_row not in (1, 2, 3):
        client_images_per_row = CLIENT_DOCS_IMAGES_PER_ROW
    client_images_per_page = client_images_per_row * client_images_per_row

    log(
        f"merge start: assets={len(asset_images)} valuation={len(valuation_images)} "
        f"client={len(client_images)} output={'disk' if output_path else 'stdout'}"
    )

    in_buf = io.BytesIO(template_bytes)
    del template_bytes
    modified: dict[str, bytes] = {}
    removed_parts: set[str] = set()
    variable_names_found: list[str] = []
    variables_occurrences_found = 0
    variables_filled = 0
    header_wraps_normalized = 0
    img_stats = {"asset": 0, "valuation": 0, "client": 0}
    preparer_stats = {
        "tableFound": 0,
        "rowsRemoved": 0,
        "preparersInserted": 0,
        "signaturesInserted": 0,
    }

    with zipfile.ZipFile(in_buf, "r") as zin:
        names = zin.namelist()

        for fname in names:
            if MERGE_PARTS_RE.match(fname):
                raw = zin.read(fname)
                for variable_name in collect_template_placeholder_names(raw):
                    if variable_name not in variable_names_found:
                        variable_names_found.append(variable_name)
                # طبّق الخطوط قبل استبدال المتغيرات حتى تبقى علامات
                # «الرقم_المرجعي» / تاريخ التقرير ظاهرة لتمييز أسفل الغلاف.
                updated = apply_report_fonts_to_part(raw)
                updated, found, filled = apply_visible_variable_values(
                    updated,
                    text_values,
                )
                if fname.lower().startswith("word/header"):
                    updated, changed_wraps = normalize_header_floating_wraps(updated)
                    header_wraps_normalized += changed_wraps
                validate_part_xml(updated, fname)
                modified[fname] = updated
                variables_occurrences_found += found
                variables_filled += filled

        if "word/fontTable.xml" in names:
            modified["word/fontTable.xml"] = ensure_tajawal_in_font_table(
                zin.read("word/fontTable.xml")
            )
        if "word/styles.xml" in names:
            modified["word/styles.xml"] = apply_tajawal_to_styles(
                zin.read("word/styles.xml")
            )
        if "word/numbering.xml" in names:
            modified["word/numbering.xml"] = apply_tajawal_to_styles(
                zin.read("word/numbering.xml")
            )

        if report_preparers_present:
            if "word/document.xml" not in names:
                raise ValueError("Invalid docx package: word/document.xml is missing")
            if "[Content_Types].xml" not in names:
                raise ValueError("Invalid docx package: [Content_Types].xml is missing")
            empty_document_relationships = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<Relationships xmlns="{PACKAGE_REL_NS}"/>'
            ).encode("utf-8")
            (
                updated_document,
                updated_document_relationships,
                updated_content_types,
                signature_parts,
                preparer_stats,
            ) = inject_report_preparers(
                modified.get("word/document.xml", zin.read("word/document.xml")),
                modified.get(
                    "word/_rels/document.xml.rels",
                    (
                        zin.read("word/_rels/document.xml.rels")
                        if "word/_rels/document.xml.rels" in names
                        else empty_document_relationships
                    ),
                ),
                modified.get(
                    "[Content_Types].xml",
                    zin.read("[Content_Types].xml"),
                ),
                report_preparers,
                set(names) | set(modified),
            )
            validate_part_xml(updated_document, "word/document.xml")
            etree.fromstring(updated_document_relationships)
            etree.fromstring(updated_content_types)
            modified["word/document.xml"] = updated_document
            modified["word/_rels/document.xml.rels"] = (
                updated_document_relationships
            )
            modified["[Content_Types].xml"] = updated_content_types
            modified.update(signature_parts)

        should_clean_mail_merge = variables_occurrences_found > 0
        if should_clean_mail_merge and "word/settings.xml" in names:
            modified["word/settings.xml"] = strip_mail_merge_settings(
                zin.read("word/settings.xml")
            )
        if should_clean_mail_merge and "word/_rels/settings.xml.rels" in names:
            modified["word/_rels/settings.xml.rels"] = strip_mail_merge_relationships(
                zin.read("word/_rels/settings.xml.rels")
            )
        if should_clean_mail_merge and "[Content_Types].xml" in names:
            modified["[Content_Types].xml"] = strip_recipient_content_type(
                modified.get(
                    "[Content_Types].xml",
                    zin.read("[Content_Types].xml"),
                )
            )
        if should_clean_mail_merge:
            removed_parts.update(
                {
                    "word/recipientData.xml",
                    "word/_rels/recipientData.xml.rels",
                }
            )

        result = write_docx_zip(zin, modified, removed_parts)

    del in_buf
    gc.collect()

    if asset_images or valuation_images or client_images:
        result, img_stats = apply_image_sections_docx_api(
            result,
            asset_images,
            valuation_images,
            client_images,
            images_per_row,
            images_per_page,
            client_images_per_row,
            client_images_per_page,
            output_path=output_path,
        )
    else:
        if output_path:
            if report_preparers_present:
                validate_docx_package(result)
            with open(output_path, "wb") as fh:
                fh.write(result)
            result = None
        else:
            validate_docx_package(result)

    log(
        json.dumps(
            {
                "variablesFound": variable_names_found,
                "variablesOccurrencesFound": variables_occurrences_found,
                "variablesFilled": variables_filled,
                "headerWrapsNormalized": header_wraps_normalized,
                "reportPreparerTableFound": preparer_stats.get("tableFound", 0),
                "reportPreparerRowsRemoved": preparer_stats.get("rowsRemoved", 0),
                "reportPreparersInserted": preparer_stats.get(
                    "preparersInserted",
                    0,
                ),
                "reportSignaturesInserted": preparer_stats.get(
                    "signaturesInserted",
                    0,
                ),
                "assetImagesInserted": img_stats.get("asset", 0),
                "valuationImagesInserted": img_stats.get("valuation", 0),
                "clientImagesInserted": img_stats.get("client", 0),
            },
            ensure_ascii=False,
        )
    )
    return result


def main() -> None:
    try:
        if len(sys.argv) > 1:
            with open(sys.argv[1], encoding="utf-8") as fh:
                payload = json.load(fh)
        else:
            payload = json.load(sys.stdin)
        result = merge_package(payload)
        if result is not None:
            sys.stdout.buffer.write(result)
        else:
            sys.stdout.write("OK\n")
    except Exception:
        log(traceback.format_exc())
        sys.exit(1)


if __name__ == "__main__":
    main()
